import os
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, timedelta, date
import tkintermapview
import requests
import locale
import csv
import sqlite3
import subprocess
import glob
import sys
from PIL import ImageTk, Image
from tkcalendar import DateEntry
import shutil
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
#from openpyxl.utils import get_column_letter
from pathlib import Path
sys.path.append(str(Path(__file__).parent / "Modüller"))
from Modüller import ajanda
from Modüller import iptv_modul
from Modüller import snake_game
from Modüller.CowsAndBulls import CowsAndBullsGame
from Modüller.AnalogSaatEmbed import AnalogSaatEmbed
from Modüller.harita import MapViewer
from Modüller.hakkinda import show_about

# HBTC Kalite Kontrol Formu oluşturma için python-docx kütüphanesi
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_BREAK
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("Uyarı: Word formu oluşturma için 'python-docx' kütüphanesi bulunamadı. 'pip install python-docx' ile kurabilirsiniz.")

# Tablolarda Türkçe alfabetik sıralama için locale kütüphanesi
try:
    locale.setlocale(locale.LC_COLLATE, 'tr_TR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_COLLATE, 'turkish')
    except locale.Error:
        print("Uyarı: Türkçe yerel ayarları (tr_TR.UTF-8 veya turkish) bulunamadı. Sıralama varsayılan şekilde yapılacak.")

ANA_DIZIN = os.path.dirname(os.path.abspath(__file__)) # Ana Dizin 
MODULES_DIR = os.path.join(ANA_DIZIN, "Modüller") # Modülleri koyduğumuz dizin 
RESOURCES_DIR = os.path.join(ANA_DIZIN, "Resources") # Programda kullanılan dosyaları koyduğumuz dizin (ikon vs)
BACKUP_DIR = os.path.join(ANA_DIZIN, "Yedeklenmis Veriler") # .csv Yedekleme dizini
EXCEL_OUTPUT_DIR = os.path.join(ANA_DIZIN, "Excel'e Aktarılanlar") # Excel çıktıları için çıktı klasörü
HBTC_FORM_OUTPUT_DIR = os.path.join(ANA_DIZIN, "HBTC Kalite Kontrol") # Word formları için çıktı klasörü
SABLONLAR_DIR = os.path.join(ANA_DIZIN, "Sablonlar") # Word ve Excel çıktıları için şablon klasörü

VERITABANI_DOSYASI = os.path.join(ANA_DIZIN, "veriler.db") # Database dosyası
HBTC_SABLON_DOSYASI = os.path.join(SABLONLAR_DIR,"HBTC_KALITE_KONTROL_FORMU.docx")
SABLON_KARSILASTIRMA_DOSYASI = os.path.join(SABLONLAR_DIR, "GLUKOMETRE_CIHAZI_KARSILASTIRMA_SONUC_FORMU.xlsx")
KALITE_KONTROL_EXCEL_SABLON_DOSYASI = os.path.join(SABLONLAR_DIR,"Kalite_Kontrol_Verileri_Sablon.xlsx")
YUZDE_SAPMA_EXCEL_SABLON_DOSYASI = os.path.join(SABLONLAR_DIR,"Yuzde_Sapma_Verileri_Sablon.xlsx")
CALENDAR_ICON_PATH = os.path.join(RESOURCES_DIR, "calendar.ico")  # Takvim simgesi
APP_ICON_PATH = os.path.join(RESOURCES_DIR, "app_icon.ico") # Uygulama simgesi

PROGRAM_AYARLARI_TABLO_ADI = "program_ayarlari"
BACKGROUND_COLOR = "#494848"

class ToolTip:
    def __init__(self, widget, text='widget info'):
        self.widget = widget
        self.text = text
        self.widget.bind("<Enter>", self.enter)
        self.widget.bind("<Leave>", self.leave)
        self.id = None
        self.tw = None # Toplevel penceresi için referans

    def enter(self, event=None):
        self.schedule()

    def leave(self, event=None):
        self.unschedule()
        self.hidetip()

    def schedule(self):
        self.unschedule()
        # Tooltip'in görünmeden önceki bekleme süresi (milisaniye)
        self.id = self.widget.after(500, self.showtip) 

    def unschedule(self):
        id_ = self.id
        self.id = None
        if id_:
            self.widget.after_cancel(id_)

    # --- BU METODU GÜNCELLEYİN ---
    def showtip(self, event=None): 
        # Eğer tooltip penceresi zaten varsa veya metin yoksa bir şey yapma
        if self.tw or not self.text:
            return

        # Widget'ın görünür olup olmadığını kontrol et
        if not self.widget.winfo_ismapped():
            self.hidetip() # Widget görünür değilse tooltip'i gizle
            return

        # Widget'ın ekran koordinatlarını al. Tooltip'i widget'ın hemen altında ve biraz sağında gösterecek şekilde ayarla
        x = self.widget.winfo_rootx() + 20 
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 1 

        self.tw = tk.Toplevel(self.widget)
        self.tw.wm_overrideredirect(True)  # Pencere kenarlıklarını kaldır
        self.tw.wm_geometry(f"+{x}+{y}")  # Pozisyonu ayarla

        label = tk.Label(self.tw, text=self.text, justify='left',
                         background="#ffffe0", relief='solid', borderwidth=1,
                         font=("tahoma", "10", "normal"))
        label.pack(ipadx=1)

    def hidetip(self):
        tw = self.tw
        self.tw = None # Referansı sıfırla
        if tw:
            tw.destroy()

class MainWindow:
    def __init__(self, master):
        self.master = master
        master.title("GLUKOMETRE TAKİP PROGRAMI")
        master.geometry("1300x700")
        self.master.configure(bg=BACKGROUND_COLOR)
        master.resizable(True, True)
        master.iconbitmap(APP_ICON_PATH)

        # Veritabanından kayıtlı pencere durumunu yükle
        is_maximized = self.program_ayari_yukle("window_maximized", "0")
        window_geometry = self.program_ayari_yukle("window_geometry", "")

        if is_maximized == "1":
            self.master.state('zoomed')  # Pencereyi maximize et
        elif window_geometry:
            try:
                self.master.geometry(window_geometry)
            except:
                pass  # Geçersiz geometri varsa varsayılanı kullan

        self.style = ttk.Style()

        for dirname in [BACKUP_DIR, EXCEL_OUTPUT_DIR, HBTC_FORM_OUTPUT_DIR, SABLONLAR_DIR]:
            if not os.path.exists(dirname):
                os.makedirs(dirname)

        self.veritabani_olustur()
        self.program_ayarlarini_yukle()
        self.style = ttk.Style()

        self.main_frame = ttk.Frame(master)
        self.main_frame.pack(fill="both", expand=True)

        self.frm_sol_panel = ttk.Frame(self.main_frame, width=270, style="SolPanel.TFrame")
        self.frm_sol_panel.pack(side="left", fill="y", padx=(10,0), pady=10)
        self.frm_sol_panel.pack_propagate(False)

        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        self.frm_glukometre_genel = ttk.LabelFrame(self.frm_sol_panel, text="Glukometre Bilgileri", style="Glukometre.TLabelframe")
        self.frm_glukometre_genel.pack(padx=5, pady=5, fill="x", expand=False)

        ttk.Label(self.frm_glukometre_genel, text="Birim/Ünite/Servis Adı:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_birim = ttk.Combobox(self.frm_glukometre_genel, state="readonly")
        self.cmb_birim.pack(fill="x", padx=5, pady=(0,5))
        self.cmb_birim.bind("<<ComboboxSelected>>", self.on_birim_cihaz_secildi)

        frm_birim_buttons = ttk.Frame(self.frm_glukometre_genel)
        frm_birim_buttons.pack(fill="x", padx=5, pady=(0,10))

        # İkonları yükle
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))            
            plus_icon_path = os.path.join(script_dir, "Resources", "plus.ico")
            minus_icon_path = os.path.join(script_dir, "Resources", "minus.ico")
            if not os.path.exists(plus_icon_path):
                raise FileNotFoundError(f"Plus ikonu bulunamadı: {plus_icon_path}")
            if not os.path.exists(minus_icon_path):
                raise FileNotFoundError(f"Minus ikonu bulunamadı: {minus_icon_path}")

            # İkonları Pillow ile aç ve PhotoImage nesnesine dönüştür
            plus_image = Image.open(plus_icon_path)
            self.plus_icon = ImageTk.PhotoImage(plus_image)
            minus_image = Image.open(minus_icon_path)
            self.minus_icon = ImageTk.PhotoImage(minus_image)

        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.plus_icon = None
            self.minus_icon = None
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.plus_icon = None
            self.minus_icon = None

        # Butonları oluştur ve ikonları ekle
        if self.plus_icon:
            btn_birim_ekle = ttk.Button(frm_birim_buttons, text="Birim Ekle", image=self.plus_icon, 
                                        compound=tk.LEFT, command=self.birim_ekle_pencere)
        else:
            btn_birim_ekle = ttk.Button(frm_birim_buttons, text="Birim Ekle", command=self.birim_ekle_pencere)
        self.btn_birim_ekle = btn_birim_ekle
        btn_birim_ekle.pack(side="left", fill="x", expand=True, padx=(0,2))
        ToolTip(self.btn_birim_ekle, "Yeni bir birim eklemek için tıklayınız.")

        if self.minus_icon:
            btn_birim_sil = ttk.Button(frm_birim_buttons, text="Birim Sil", image=self.minus_icon, 
                                       compound=tk.LEFT, command=self.birim_sil)
        else:
            btn_birim_sil = ttk.Button(frm_birim_buttons, text="Birim Sil", command=self.birim_sil)
        self.btn_birim_sil = btn_birim_sil
        btn_birim_sil.pack(side="left", fill="x", expand=True, padx=(2,0))
        ToolTip(self.btn_birim_sil, "Seçili birimi silmek için tıklayınız.")

        ttk.Label(self.frm_glukometre_genel, text="Cihaz Tipi - Marka:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_device_type = ttk.Combobox(self.frm_glukometre_genel, state="readonly")
        self.cmb_device_type.pack(fill="x", padx=5, pady=(0,5))
        self.cmb_device_type.bind("<<ComboboxSelected>>", self.on_device_type_selected)

        # device_type_buttons Frame'ini ekle
        frm_device_type_buttons = ttk.Frame(self.frm_glukometre_genel)
        frm_device_type_buttons.pack(fill="x", padx=5, pady=(0,10))

        # Butonları oluştur ve ikonları ekle
        if self.plus_icon:
            btn_device_type_ekle = ttk.Button(frm_device_type_buttons, 
                                            text="Cihaz Markası Ekle", 
                                            image=self.plus_icon, 
                                            compound=tk.LEFT, 
                                            command=self.cihaz_markasi_ekle_pencere)
        else:
            btn_device_type_ekle = ttk.Button(frm_device_type_buttons, 
                                            text="Cihaz Markası Ekle", 
                                            command=self.cihaz_markasi_ekle_pencere)
        self.btn_device_type_ekle = btn_device_type_ekle
        btn_device_type_ekle.pack(side="left", fill="x", expand=True, padx=(0,2))
        ToolTip(self.btn_device_type_ekle, "Yeni bir cihaz markası eklemek için tıklayınız.")

        if self.minus_icon:
            btn_device_type_sil = ttk.Button(frm_device_type_buttons, 
                                            text="Cihaz Markası Sil", 
                                            image=self.minus_icon, 
                                            compound=tk.LEFT, 
                                            command=self.cihaz_markasi_sil)
        else:
            btn_device_type_sil = ttk.Button(frm_device_type_buttons, 
                                            text="Cihaz Markası Sil", 
                                            command=self.cihaz_markasi_sil)
        self.btn_device_type_sil = btn_device_type_sil
        btn_device_type_sil.pack(side="left", fill="x", expand=True, padx=(2,0))
        ToolTip(self.btn_device_type_sil, "Seçili cihaz markasını silmek için tıklayınız.")

        frm_seri_no_container = ttk.Frame(self.frm_glukometre_genel)
        frm_seri_no_container.pack(fill="x", padx=0, pady=0)
        frm_seri_no_sol = ttk.Frame(frm_seri_no_container)
        frm_seri_no_sol.pack(side="left", fill="x", expand=True, padx=(5,2), pady=(0,5))
        
        ttk.Label(frm_seri_no_sol, text="Cihaz Seri No:").pack(fill="x")
        self.cmb_device_serial = ttk.Combobox(frm_seri_no_sol, state="readonly")
        self.cmb_device_serial.pack(fill="x")

        frm_seri_no_sag = ttk.Frame(frm_seri_no_container)
        frm_seri_no_sag.pack(side="left", fill="x", expand=True, padx=(2,5), pady=(0,5))
        ttk.Label(frm_seri_no_sag, text="Son 4 Hane:").pack(fill="x")

        self.cmb_son4hane = ttk.Combobox(frm_seri_no_sag)
        self.cmb_son4hane.pack(fill="x")
        self.cmb_son4hane.bind("<<ComboboxSelected>>", self.on_son4hane_changed)
        self.cmb_son4hane.bind("<FocusOut>", self.on_son4hane_changed)
        self.cmb_son4hane.bind("<KeyRelease>", self.validate_son4hane_input)

        # --- Analog Saat Alanı ---
        # Analog Saat Alanı
        clock_widget_actual_size = 220  # frm_sol_panel'in genişliği 270. Saat için uygun bir boyut seçelim. Saatin kaplayacağı gerçek piksel boyutu (kare)
                                        # Bu değeri frm_sol_panel genişliğine ve istediğiniz görünüme göre ayarlayabilirsiniz.
        sol_panel_bg = self.style.lookup('SolPanel.TFrame', 'background')
        if not sol_panel_bg:
            sol_panel_bg = "white"

        self.analog_saat_canvas = tk.Canvas(self.frm_sol_panel,
                                            width=clock_widget_actual_size,
                                            height=clock_widget_actual_size,
                                            bg=sol_panel_bg,
                                            highlightthickness=0, cursor="hand2")
        self.analog_saat_canvas.pack(side="top", pady=(5, 5))

        try:
            self.analog_saat_instance = AnalogSaatEmbed(self.analog_saat_canvas, size=clock_widget_actual_size, canvas_bg_color=sol_panel_bg)
            self.analog_saat_instance.start()
        except Exception as e:
            print(f"Analog saat başlatılamadı: {e}")
            self.analog_saat_instance = None

        # Analog Saate Tooltip Ekleme
        ToolTip(self.analog_saat_canvas, "Alarm, Geri Sayım ve Kronometre aracını açmak için tıklayın.")
        
        # Analog Saate Tıklama Olayı Ekleme
        self.analog_saat_canvas.bind("<Button-1>", lambda event: self.launch_countdown_alarm())

        self.lbl_tarih = ttk.Label(self.frm_sol_panel, text=self.get_turkish_datetime_str(datetime.now()), anchor="center", font=("Arial", 14, "bold"), foreground="#0F2E4C", cursor="hand2")
        self.lbl_tarih.pack(pady=(0, 15))
        self.lbl_tarih.bind("<Button-1>", lambda e: self.launch_calendar())
        ToolTip(self.lbl_tarih, "Ajandayı açmak için tıklayın.")

        self.frm_radyo = ttk.LabelFrame(self.frm_sol_panel, text="Radyo", style="Radyo.TLabelframe")
        self.frm_radyo.pack(side="bottom", fill="x")

        self.cmb_radyo = ttk.Combobox(self.frm_radyo, state="readonly")
        self.radio_station_names, self.radio_station_map = self.radyo_istasyonlar_comboboxa_yukle()
        ttk.Label(self.frm_radyo, text="Radyo İstasyonu:").pack(fill="x", padx=5, pady=(5,0))
        self.cmb_radyo['values'] = self.radio_station_names

        last_radio_station = self.program_ayari_yukle("last_radio_station")
        if last_radio_station and last_radio_station in self.radio_station_names:
            self.cmb_radyo.set(last_radio_station)
        elif self.radio_station_names:
            self.cmb_radyo.current(0)

        self.radyo_istasyonlar_comboboxa_yukle()
        self.cmb_radyo.pack(fill="x", padx=5, pady=(0,5))
        last_radio_station = self.program_ayari_yukle("last_radio_station")
        if last_radio_station and last_radio_station in self.radio_station_names:
            self.cmb_radyo.set(last_radio_station)
        elif self.radio_station_names:
            self.cmb_radyo.current(0)

        frm_radyo_controls = ttk.Frame(self.frm_radyo)
        frm_radyo_controls.pack(fill="x", padx=5, pady=(0,10))

        # Play ve Stop ikonlarını yükle
        try:
            # Betiğin bulunduğu dizini al
            script_dir = os.path.dirname(os.path.abspath(__file__))
            play_icon_path = os.path.join(script_dir, "Resources", "play.ico")
            stop_icon_path = os.path.join(script_dir, "Resources", "stop.ico")
            if not os.path.exists(play_icon_path):
                raise FileNotFoundError(f"Play ikonu bulunamadı: {play_icon_path}")
            if not os.path.exists(stop_icon_path):
                raise FileNotFoundError(f"Stop ikonu bulunamadı: {stop_icon_path}")
            # İkonları Pillow ile aç ve PhotoImage nesnesine dönüştür
            play_image = Image.open(play_icon_path)
            self.play_icon = ImageTk.PhotoImage(play_image)
            stop_image = Image.open(stop_icon_path)
            self.stop_icon = ImageTk.PhotoImage(stop_image)

        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.play_icon = None
            self.stop_icon = None
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.play_icon = None
            self.stop_icon = None

        # Butonları oluştur ve ikonları ekle
        if self.play_icon:
            self.btn_radyo_play = ttk.Button(frm_radyo_controls, text="Play", image=self.play_icon, compound=tk.LEFT, command=self.play_radio_command)
        else:
            self.btn_radyo_play = ttk.Button(frm_radyo_controls, text="Play", command=self.play_radio_command)
        self.btn_radyo_play.grid(row=0, column=0, sticky="ew", padx=(0,1))

        if self.stop_icon:
            self.btn_radyo_stop = ttk.Button(frm_radyo_controls, text="Stop", image=self.stop_icon, compound=tk.LEFT, command=self.stop_radio)
        else:
            self.btn_radyo_stop = ttk.Button(frm_radyo_controls, text="Stop", command=self.stop_radio)
        self.btn_radyo_stop.grid(row=0, column=1, sticky="ew", padx=(1,0))

        self.radio_volume = tk.IntVar(value=50)
        last_volume = self.program_ayari_yukle("last_radio_volume", "50")
        try:
            self.radio_volume.set(int(last_volume))
        except ValueError:
            self.radio_volume.set(50)

        self.volume_slider = ttk.Scale(frm_radyo_controls, from_=0, to=100, orient=tk.HORIZONTAL, 
                                       variable=self.radio_volume, command=self.on_volume_change)
        self.volume_slider.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(5,0))
        
        # "Sesi Kıs" butonu
        try:
            # Betiğin bulunduğu dizini al
            script_dir = os.path.dirname(os.path.abspath(__file__))
            mute_icon_path = os.path.join(script_dir, "Resources", "mute.ico")
            loud_icon_path = os.path.join(script_dir, "Resources", "loud.ico")

            if not os.path.exists(mute_icon_path):
                raise FileNotFoundError(f"Mute ikonu bulunamadı: {mute_icon_path}")
            if not os.path.exists(loud_icon_path):
                raise FileNotFoundError(f"Loud ikonu bulunamadı: {loud_icon_path}")

            self.mute_icon = ImageTk.PhotoImage(Image.open(mute_icon_path))
            self.loud_icon = ImageTk.PhotoImage(Image.open(loud_icon_path))

            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", image=self.mute_icon, 
                                             compound=tk.LEFT, command=self.toggle_mute_sound)
        except FileNotFoundError as fnf_error:
            print(f"İkon dosyası hatası: {fnf_error}. 'Resources' klasörünün doğru yerde olduğundan ve ikon dosyalarının mevcut olduğundan emin olun.")
            self.mute_icon = None
            self.loud_icon = None
            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", command=self.toggle_mute_sound)
        except Exception as e: # Diğer olası Pillow veya Tkinter hatalarını yakala
            print(f"İkon yükleme sırasında genel hata: {e}")
            self.mute_icon = None
            self.loud_icon = None
            self.btn_mute_sound = ttk.Button(frm_radyo_controls, text="Sesi Kıs", command=self.toggle_mute_sound)

        self.btn_mute_sound.grid(row=1, column=2, sticky="ew", padx=(5,0), pady=(5,0))
        self.lbl_volume_percent = ttk.Label(frm_radyo_controls, text="")
        self.lbl_volume_percent.grid(row=2, column=0, columnspan=3, sticky="n")

        frm_radyo_controls.columnconfigure(0, weight=1)
        frm_radyo_controls.columnconfigure(1, weight=1)
        frm_radyo_controls.columnconfigure(2, weight=0)

        self.radio_process = None
        self.update_radio_button_states()
        self.cmb_radyo.bind("<<ComboboxSelected>>", self.comboboxtan_radyo_degistir)

        # Mute durumunu takip etmek için bir değişken
        self.is_muted = False
        self.previous_volume = self.radio_volume.get() # Başlangıçta mevcut ses seviyesini sakla

        self.status_bar = tk.Label(self.master, text="", relief="sunken", anchor="w", font=("Arial", 10, "bold"))
        self.status_bar.pack(side="bottom", fill="x")
        self.veritabanindan_verileri_cek()

        # Combobox seçim değişikliklerini ayarlara kaydet
        self.cmb_device_type.bind("<<ComboboxSelected>>", self.on_device_type_selected)
        self.cmb_device_serial.bind("<<ComboboxSelected>>", self.on_device_serial_selected)
        self.on_birim_cihaz_secildi()

        self.measurement_no_kalite = 1
        self.measurement_no_yuzde = 1
        self.editing_entry = None
        self.l_entry_tooltips = {}
        self.tables_cleared_this_session = False

        self.create_tabs()
        self.create_menu()
        self.otomatik_yedek_yukle()
        self.statusbar_guncelle()
        master.protocol("WM_DELETE_WINDOW", self.kaydet_ve_cikis_yap)

        # Cihaz kaydı çakışma kontrolünü tablardaki textbox kutularına odaklanınca yap
        self.txt_l1.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_l2.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_l3.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_glukometre_yuzde.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.txt_lab_yuzde.bind("<FocusIn>", self.kontrol_cihaz_kayit_cakisma)
        self.cakisma_uyarildi = False  # Uyarı kutusunun sadece 1 kez çıkması için gösterildi mi kontrolü
        self.iptv_channels = []  # IPTV kanallarını saklamak için liste

        self.map_viewer = MapViewer(self.master)

    def launch_countdown_alarm(self):
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            countdown_alarm_path = os.path.join(script_dir, "Modüller", "Countdown_Alarm.py")
            if not os.path.exists(countdown_alarm_path):
                messagebox.showerror("Hata", f"'Countdown_Alarm.py' dosyası bulunamadı:\n{countdown_alarm_path}", parent=self.master)
                return
            subprocess.Popen([sys.executable, countdown_alarm_path])
        except Exception as e:
            messagebox.showerror("Hata", f"Countdown Alarm uygulaması başlatılamadı:\n{e}", parent=self.master)


    def get_turkish_datetime_str(self, dt_object):
        days_tr = ["Pazartesi", "Salı", "Çarşamba", "Perşembe", "Cuma", "Cumartesi", "Pazar"]
        months_tr_short = ["Oca", "Şub", "Mar", "Nis", "May", "Haz",
                        "Tem", "Ağu", "Eyl", "Eki", "Kas", "Ara"]

        day_name = days_tr[dt_object.weekday()]
        month_name = months_tr_short[dt_object.month - 1]
        
        return f"{dt_object.day} {month_name} {dt_object.year} {day_name}"

    def create_tabs(self):
        self.tab1 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text="Kalite Kontrol Ölçümleri")
        self.tab1.grid_columnconfigure(0, weight=1)
        self.tab1.grid_rowconfigure(2, weight=1)

        lbl_tablo_baslik_kalite = ttk.Label(self.tab1, text="KALİTE KONTROL ÖLÇÜMLERİ", font=("Tahoma", 14, "bold"), foreground="#09767E")
        lbl_tablo_baslik_kalite.grid(row=0, column=0, pady=(10,5), sticky="n")
        self.tab1.grid_rowconfigure(0, weight=0)

        frm_kalite_input = ttk.LabelFrame(self.tab1, text="Kalite Kontrol Değer Girişi", style="KaliteInput.TLabelframe")
        frm_kalite_input.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        vcmd_l1 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 36, 108)), '%P', '%W')
        vcmd_l2 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 144, 216)), '%P', '%W')
        vcmd_l3 = (self.master.register(lambda P, w_name: self.validate_l_entry(P, w_name, 252, 396)), '%P', '%W')
        
        lbl_l1 = ttk.Label(frm_kalite_input, text="L1 Ölçümü:", font= ("Arial", 12, "bold"), foreground="#FFEE00")
        lbl_l1.grid(row=0, column=0, padx=5, pady=5)
        self.txt_l1 = ttk.Entry(frm_kalite_input, width=4, font= ("Arial", 12, "bold"), name="l1_entry", justify='center', 
                                validate="key", validatecommand=vcmd_l1)
        self.txt_l1.grid(row=0, column=1, padx=5, pady=5)
        self.l_entry_tooltips["l1_entry"] = ToolTip(self.txt_l1, "Seviye 1 (L1) Değeri 36-108 arası olmalı")

        lbl_l2 = ttk.Label(frm_kalite_input, text="L2 Ölçümü:", font= ("Arial", 12, "bold"), foreground="#0000FF")
        lbl_l2.grid(row=0, column=2, padx=5, pady=5)
        self.txt_l2 = ttk.Entry(frm_kalite_input, width=4, font= ("Arial", 12, "bold"), name="l2_entry", justify='center', 
                                validate="key", validatecommand=vcmd_l2)
        self.txt_l2.grid(row=0, column=3, padx=5, pady=5)
        self.l_entry_tooltips["l2_entry"] = ToolTip(self.txt_l2, "Seviye 2 (L2) Değeri  144-216 arası olmalı")

        lbl_l3 = ttk.Label(frm_kalite_input, text="L3 Ölçümü:", font= ("Arial", 12, "bold"), foreground="#FF0000")
        lbl_l3.grid(row=0, column=4, padx=5, pady=5)
        self.txt_l3 = ttk.Entry(frm_kalite_input, width=4, font= ("Arial", 12, "bold"), name="l3_entry", justify='center', 
                                validate="key", validatecommand=vcmd_l3)
        self.txt_l3.grid(row=0, column=5, padx=5, pady=5)
        self.l_entry_tooltips["l3_entry"] = ToolTip(self.txt_l3, "Seviye 3 (L3) Değeri 252-396 arası olmalı")

        ttk.Button(frm_kalite_input, text="Tabloya Aktar", command=self.tabloya_aktar_kalite).grid(row=0, column=6, padx=20, pady=5)

        columns_kalite = (
            "No", "Tarih", "Cihaz Tipi - Marka", "Cihaz Seri No", "L1", "L2", "L3",
            "Birim/Ünite/Servis Adı", "Bir Sonraki Gelinecek Tarih"
        )
        self.tree_kalite = ttk.Treeview(self.tab1, columns=columns_kalite, show="headings", selectmode="extended")
        self.tree_kalite.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

        vsb_kalite = ttk.Scrollbar(self.tab1, orient="vertical", command=self.tree_kalite.yview)
        vsb_kalite.grid(row=2, column=1, sticky='ns', pady=(10,10))
        hsb_kalite = ttk.Scrollbar(self.tab1, orient="horizontal", command=self.tree_kalite.xview)
        hsb_kalite.grid(row=3, column=0, sticky='ew', padx=(10,10))
        self.tree_kalite.configure(yscrollcommand=vsb_kalite.set, xscrollcommand=hsb_kalite.set)

        widths_kalite = [20, 50, 150, 80, 30, 30, 30, 200, 140] # Sütun genişlikleri
        for col, w in zip(columns_kalite, widths_kalite):
            self.tree_kalite.heading(col, text=col, command=lambda c=col: self.treeview_sort_column(self.tree_kalite, c, False))
            self.tree_kalite.column(col, width=w, anchor=tk.CENTER, minwidth=w)

        self.tree_kalite.bind("<Double-1>", lambda event: self.on_double_click(event, self.tree_kalite))
        self.tree_kalite.bind("<Button-3>", lambda event: self.show_context_menu(event, self.tree_kalite))
        self.tree_kalite.bind("<Delete>", lambda event: self.satir_sil(self.tree_kalite, True))

        self.context_menu_kalite = tk.Menu(self.tab1, tearoff=0)
        self.context_menu_kalite.add_command(label="Satırı Sil", command=lambda: self.satir_sil(self.tree_kalite, True))

        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab2, text="Yüzde Sapma Ölçümleri")
        self.tab2.grid_columnconfigure(0, weight=1)
        self.tab2.grid_rowconfigure(2, weight=1)

        lbl_tablo_baslik_yuzde = ttk.Label(self.tab2, text="YÜZDE SAPMA ÖLÇÜMLERİ", font=("Tahoma", 14, "bold"), foreground="#09767E")
        lbl_tablo_baslik_yuzde.grid(row=0, column=0, pady=(10,5), sticky="n")
        self.tab2.grid_rowconfigure(0, weight=0)

        frm_yuzde_input = ttk.LabelFrame(self.tab2, text="Yüzde Sapma Değer Girişi", style="YuzdeInput.TLabelframe")
        frm_yuzde_input.grid(row=1, column=0, padx=5, pady=5, sticky="ew")

        ttk.Label(frm_yuzde_input, text="Hasta Ad Soyad:", font=("Arial", 10, "bold"), foreground="#000000").grid(row=0, column=0, padx=5, pady=5)
        
        self.txt_hasta_ad_soyad = ttk.Entry(frm_yuzde_input, width=35, font=("Arial", 10, "bold"), justify='left')
        self.txt_hasta_ad_soyad.grid(row=0, column=1, padx=5, pady=5)
        ToolTip(self.txt_hasta_ad_soyad, "Gelen kan numunesinin ait olduğu hastanın adı ve soyadı.")

        def on_hasta_ad_soyad_key_release(event): #Hasta ad soyad daima büyük harf yazılır
            widget = event.widget
            current_text = widget.get()
            cursor_pos = widget.index(tk.INSERT)
            widget.delete(0, tk.END)
            widget.insert(0, current_text.upper())
            widget.icursor(cursor_pos)

        self.txt_hasta_ad_soyad.bind("<KeyRelease>", on_hasta_ad_soyad_key_release)

        ttk.Label(frm_yuzde_input, text="Glukometre Sonucu (mg/dl):", font=("Arial", 10, "bold"), foreground="#000000").grid(row=0, column=2, padx=10, pady=5)   
        
        self.txt_glukometre_yuzde = ttk.Entry(frm_yuzde_input, width=4, font= ("Arial", 10, "bold"), justify='center')
        self.txt_glukometre_yuzde.grid(row=0, column=3, padx=5, pady=5)
        ToolTip(self.txt_glukometre_yuzde, "Gelen kan numunesinin Glukometre cihazında ölçtüğünüz ölçüm değerini giriniz.")    
        
        ttk.Label(frm_yuzde_input, text="Otoanalizör Sonucu (mg/dl):", font= ("Arial", 10, "bold"), foreground="#000000").grid(row=0, column=4, padx=10, pady=5)
        self.txt_lab_yuzde = ttk.Entry(frm_yuzde_input, width=4, font= ("Arial", 10, "bold"), justify='center')
        self.txt_lab_yuzde.grid(row=0, column=5, padx=5, pady=5)
        ToolTip(self.txt_lab_yuzde, "Gelen kan numunesinin Laboratuar Otoanalizör cihazında ölçülen ölçüm değerini giriniz.")    

        ttk.Button(frm_yuzde_input, text="Hesapla ve Tabloya Aktar", command=self.yuzde_sapma_hesapla_ve_aktar).grid(row=0, column=6, padx=5, pady=5)

        columns_yuzde = ("No", "Tarih", "Cihaz Marka", "Cihaz Seri No", "Birim/Ünite/Servis Adı", "Hasta Ad Soyad", "Glukometre Sonucu", "Oto analizör Sonucu", "% Sapma Oranı", "Değerlendirme Sonucu", "Bir Sonraki Gelinecek Tarih")
        self.tree_yuzde = ttk.Treeview(self.tab2, columns=columns_yuzde, show="headings", selectmode="extended")
        self.tree_yuzde.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

        vsb_yuzde = ttk.Scrollbar(self.tab2, orient="vertical", command=self.tree_yuzde.yview)
        vsb_yuzde.grid(row=2, column=1, sticky='ns', pady=(10,10))
        hsb_yuzde = ttk.Scrollbar(self.tab2, orient="horizontal", command=self.tree_yuzde.xview)
        hsb_yuzde.grid(row=3, column=0, sticky='ew', padx=(10,10))
        self.tree_yuzde.configure(yscrollcommand=vsb_yuzde.set, xscrollcommand=hsb_yuzde.set)
        self.tree_yuzde.tag_configure('high_deviation_tree', background='red', foreground='white')

        widths_yuzde = [30, 65, 85, 110, 170, 170, 50, 50, 60, 120, 120]
        for col, w in zip(columns_yuzde, widths_yuzde):
            self.tree_yuzde.heading(col, text=col, command=lambda c=col: self.treeview_sort_column(self.tree_yuzde, c, False))
            self.tree_yuzde.column(col, width=w, anchor=tk.W, stretch=False, minwidth=w)

        self.tree_yuzde.bind("<Double-1>", lambda event: self.on_double_click(event, self.tree_yuzde))
        self.tree_yuzde.bind("<Button-3>", lambda event: self.show_context_menu(event, self.tree_yuzde))
        self.tree_yuzde.bind("<Delete>", lambda event: self.satir_sil(self.tree_yuzde, False))

        self.context_menu_yuzde = tk.Menu(self.tab2, tearoff=0)
        self.context_menu_yuzde.add_command(label="Satırı Sil", command=lambda: self.satir_sil(self.tree_yuzde, False))


    def program_ayarlarini_yukle(self):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS {PROGRAM_AYARLARI_TABLO_ADI} (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)
        conn.commit()
        conn.close()

    def program_ayari_kaydet(self, key, value):
        try:
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute(f"INSERT OR REPLACE INTO {PROGRAM_AYARLARI_TABLO_ADI} (key, value) VALUES (?, ?)", (key, value))
            conn.commit()
            conn.close()
        except sqlite3.Error as e:
            print(f"DB Ayar kaydetme hatası ({key}): {e}")

    def program_ayari_yukle(self, key, default=None):
        try:
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute(f"SELECT value FROM {PROGRAM_AYARLARI_TABLO_ADI} WHERE key = ?", (key,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else default
        except sqlite3.Error as e:
            print(f"DB Ayar yükleme hatası ({key}): {e}")
            return default

    def comboboxtan_radyo_degistir(self, event=None): # Comboboxtan başka bir radyo seçilince        
        if self.radio_process and self.radio_process.poll() is None: # Radyo çalıyorsa önce durdur
            self.stop_radio()           
            self.master.after(200, self.play_radio_command) # ffplay in durmasını bekle, sonra mevcut ses seviyesi ile yeni radyo çal 

    def play_radio_command(self):
        self.play_radio()

    def play_radio(self, volume_level=None):
        selected_name = self.cmb_radyo.get()
        if not selected_name or selected_name not in self.radio_station_map:
            messagebox.showwarning("Radyo", "Lütfen bir radyo istasyonu seçin!", parent=self.master)
            return
        url = self.radio_station_map[selected_name]

        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()

        actual_volume = volume_level if volume_level is not None else self.radio_volume.get()

        try: # Önce programın çalıştığı dizindeki ffmpeg klasöründe ffplay.exe'yi ara           
            script_dir = os.path.dirname(os.path.abspath(__file__))
            ffmpeg_dir = os.path.join(script_dir, "ffmpeg")
            ffplay_path = os.path.join(ffmpeg_dir, "ffplay.exe")
            
            if not os.path.exists(ffplay_path):                
                ffplay_path = "ffplay" # Eğer yerel klasörde yoksa, sistem PATH'ında ara

            startupinfo = None; creation_flags = 0
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = subprocess.SW_HIDE
                creation_flags = subprocess.CREATE_NO_WINDOW
        
            cmd = [ffplay_path, "-nodisp", "-autoexit", "-loglevel", "error", "-volume", str(actual_volume), url]
            self.radio_process = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                                                  startupinfo=startupinfo, creationflags=creation_flags)
            self.update_radio_button_states(playing=True)
            self.start_marquee(selected_name)
        except FileNotFoundError:
            messagebox.showerror("Radyo", 
                           "ffplay bulunamadı!\n"
                           f"Lütfen ffmpeg klasörünü program dizinine ({ffmpeg_dir}) kopyalayın\n"
                           "veya sistem PATH'ına ffmpeg ekleyin.", 
                           parent=self.master)
            self.radio_process = None; self.update_radio_button_states(playing=False)
            self.lbl_volume_percent.config(text="ffplay bulunamadı!") # Hata mesajını label'da göster
        except Exception as e:
            messagebox.showerror("Radyo", f"Radyo başlatılamadı: {e}", parent=self.master)
            self.radio_process = None; self.update_radio_button_states(playing=False)
            self.lbl_volume_percent.config(text=f"Radyo başlatılamadı: {e}") # Hata mesajını label'da göster

    def stop_radio(self):
        if self.radio_process and self.radio_process.poll() is None:
            try: self.radio_process.terminate()
            except Exception as e: print(f"Radyo durdurulurken hata: {e}")
            finally: self.radio_process = None
        self.update_radio_button_states(playing=False)
        self.stop_marquee()
        self.lbl_volume_percent.config(text="")

    def on_volume_change(self, value_str):
        volume_level = int(float(value_str))
        self.radio_volume.set(volume_level)
        self.lbl_volume_percent.config(text=f"Ses Seviyesi: {volume_level}%")
        self.program_ayari_kaydet("last_radio_volume", str(volume_level))
        if self.radio_process and self.radio_process.poll() is None: # Radyo çalıyorsa önce durdur
            self.stop_radio()
            # ffplay in durmasını bekle, sonra yeni ses seviyesi ile tekrar çal 
            self.master.after(200, lambda: self.play_radio(volume_level=volume_level))        
        else:
            self.lbl_volume_percent.config(text=f"Ses Seviyesi: {volume_level}%") # Sadece ses seviyesini göster
            
    def toggle_mute_sound(self):
        if not hasattr(self, 'loud_icon') or not hasattr(self, 'mute_icon'):
            messagebox.showerror("Hata", "Loud/Mute ikonları yüklenemedi!")
            return

        if not self.is_muted:
            self.previous_volume = self.radio_volume.get()
            self.radio_volume.set(0)
            self.volume_slider.set(0)
            self.lbl_volume_percent.config(text=f"Ses Seviyesi: 0%")
            self.program_ayari_kaydet("last_radio_volume", "0")

            # Butonun görünümünü değiştir
            self.btn_mute_sound.config(image=self.loud_icon, text="Sesi Aç")
        else:
            # Sesi aç
            self.radio_volume.set(self.previous_volume)
            self.volume_slider.set(self.previous_volume)
            self.lbl_volume_percent.config(text=f"Ses Seviyesi: {self.previous_volume}%")
            self.program_ayari_kaydet("last_radio_volume", str(self.previous_volume))

            # Butonun görünümünü değiştir
            self.btn_mute_sound.config(image=self.mute_icon, text="Sesi Kıs")

        self.is_muted = not self.is_muted
        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()
            self.master.after(200, lambda: self.play_radio(volume_level=self.radio_volume.get()))
        else: # Radyo çalmıyorsa, sadece ses seviyesini sıfırla
            self.update_radio_button_states(playing=False)

    def mute_sound(self):
        self.radio_volume.set(0)
        self.volume_slider.set(0) # Slider'ı da güncelle
        self.lbl_volume_percent.config(text=f"Ses Seviyesi: 0%")
        self.program_ayari_kaydet("last_radio_volume", "0")
        if self.radio_process and self.radio_process.poll() is None:
            self.stop_radio()
        else: # Radyo çalmıyorsa, sadece ses seviyesini sıfırla
            self.update_radio_button_states(playing=False)

    def update_radio_button_states(self, playing=False):
        if playing:
            self.btn_radyo_play.config(state=tk.DISABLED)
            self.btn_radyo_stop.config(state=tk.NORMAL)
            self.btn_mute_sound.config(state=tk.NORMAL)
        else:
            self.btn_radyo_play.config(state=tk.NORMAL)
            self.btn_radyo_stop.config(state=tk.DISABLED)
            self.btn_mute_sound.config(state=tk.DISABLED)

    def start_marquee(self, station_name):
        self.marquee_text = f"..:: NOW PLAYING - {station_name} ::.."
        self.marquee_pos = 0
        self.marquee_update()

    def stop_marquee(self):
        if hasattr(self, 'marquee_job') and self.marquee_job is not None:
            self.master.after_cancel(self.marquee_job)
            self.marquee_job = None

    def marquee_update(self):
        if self.radio_process and self.radio_process.poll() is None:
            displayed_text = self.marquee_text[self.marquee_pos:] + self.marquee_text[:self.marquee_pos]
            self.lbl_volume_percent.config(text=displayed_text)
            self.marquee_pos = (self.marquee_pos + 1) % len(self.marquee_text)
            self.marquee_job = self.master.after(200, self.marquee_update)

    def statusbar_guncelle(self):
        # Kayıtlı cihaz sayısını veritabanından çek
        kayitli_cihaz_sayisi = 0
        try:
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM cihaz_kayitlari")
            result = cursor.fetchone()
            if result:
                kayitli_cihaz_sayisi = result[0]
            conn.close()
        except sqlite3.Error as e:
            print(f"DB Kayıtlı cihaz sayısı alınırken hata: {e}")

        kalite_count = len(self.tree_kalite.get_children()) if hasattr(self, 'tree_kalite') else 0
        yuzde_count = len(self.tree_yuzde.get_children()) if hasattr(self, 'tree_yuzde') else 0
        status_text = f"Envantere Kayıtlı Cihaz Sayısı: {kayitli_cihaz_sayisi}   |   Kalite Kontrol Ölçümleri: {kalite_count}   |   Yüzde Sapma Ölçümleri: {yuzde_count}"
        self.status_bar.config(text=status_text)

    def veritabani_olustur(self):
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()

            # Birimler Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS birimler (
                    birim_adi TEXT PRIMARY KEY
                )
            """)

            # Cihaz Tipleri Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS cihaz_tipleri (
                    cihaz_tipi TEXT PRIMARY KEY
                )
            """)

            # Cihaz Serileri Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS cihaz_serileri (
                    cihaz_seri TEXT PRIMARY KEY,
                    cihaz_tipi TEXT NOT NULL
                )
            """)

            # Mevcut cihaz_serileri tablosuna cihaz_tipi sütununu ekleme (eski versiyon uyumluluğu için)
            try:
                cursor.execute("ALTER TABLE cihaz_serileri ADD COLUMN cihaz_tipi TEXT NOT NULL DEFAULT 'Bilinmiyor'") # DEFAULT eklendi
            except sqlite3.OperationalError:
                pass  # Sütun zaten var veya başka bir hata (örn: NOT NULL constraint default olmadan)

            # Radyolar Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS radyolar (
                    radyo_adi TEXT PRIMARY KEY,
                    radyo_url TEXT NOT NULL
                )
            """)

            # Cihaz Kayıtları Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS cihaz_kayitlari (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    birim_adi TEXT NOT NULL,
                    cihaz_tipi TEXT NOT NULL,
                    cihaz_seri TEXT NOT NULL,
                    son_4_hane TEXT NOT NULL,
                    uretim_yili TEXT,
                    satin_alma_tarihi TEXT,
                    temsilci_firma TEXT,
                    UNIQUE (cihaz_tipi, cihaz_seri, son_4_hane)
                )
            """)

            # Üretim Tarihleri Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS uretim_tarihleri (
                    uretim_yili TEXT PRIMARY KEY
                )
            """)

            # Satın Alma Tarihleri Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS satin_alma_tarihleri (
                    satin_alma_tarihi TEXT PRIMARY KEY
                )
            """)

            # Temsilci Firmalar Tablosu
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS temsilci_firmalar (
                    temsilci_firma_adi TEXT PRIMARY KEY
                )
            """)

            conn.commit()
            conn.close()

    def veritabanindan_verileri_cek(self):
        self.birimleri_comboboxa_yukle()
        self.cihaz_tiplerini_comboboxa_yukle()
        self.cihaz_seri_no_lar_comboboxa_yukle()
        self.radyo_istasyonlar_comboboxa_yukle()
        self.initialize_agenda_module()

    def otomatik_yedek_yukle(self):# Başlangıçta Yedekleme klasöründen en son yedek dosyalarını otomatik yükle
        self.master.update_idletasks()
        try:
            load_kalite_files, load_yuzde_files = [], []
            all_kalite_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Kalite_Kontrol_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
            if all_kalite_files: load_kalite_files = [max(all_kalite_files, key=os.path.getctime)]
            all_yuzde_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Yuzde_Sapma_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
            if all_yuzde_files: load_yuzde_files = [max(all_yuzde_files, key=os.path.getctime)]

            loaded_kalite, loaded_yuzde = False, False
            for item in self.tree_kalite.get_children(): self.tree_kalite.delete(item)
            for item in self.tree_yuzde.get_children(): self.tree_yuzde.delete(item)
            self.measurement_no_kalite, self.measurement_no_yuzde = 1, 1

            expected_kalite_headers = self.tree_kalite['columns']
            for file_path in load_kalite_files:
                try:
                    with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                        reader = csv.reader(f, delimiter=';')
                        try: headers = next(reader)
                        except StopIteration: continue
                        if tuple(headers) == expected_kalite_headers:
                            file_has_data = False
                            for row in reader:
                                if len(row) == len(expected_kalite_headers): self.tree_kalite.insert("", "end", values=row); file_has_data = True
                            if file_has_data: loaded_kalite = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

            expected_yuzde_headers = self.tree_yuzde['columns']
            sapma_str_idx = list(expected_yuzde_headers).index("% Sapma Oranı") if "% Sapma Oranı" in expected_yuzde_headers else -1
            for file_path in load_yuzde_files:
                try:
                    with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                        reader = csv.reader(f, delimiter=';')
                        try: headers = next(reader)
                        except StopIteration: continue
                        if tuple(headers) == expected_yuzde_headers:
                            file_has_data = False
                            for i, row in enumerate(reader, start=1):
                                if len(row) == len(expected_yuzde_headers):
                                    tags_to_apply = ()
                                    if sapma_str_idx != -1:
                                        try:
                                            sapma_val = float(row[sapma_str_idx].replace('%','').strip())
                                            if sapma_val > 9.99: tags_to_apply = ('high_deviation_tree',)
                                        except: pass # Ignore tagging errors for malformed data
                                    self.tree_yuzde.insert("", "end", values=row, tags=tags_to_apply)
                                    file_has_data = True
                            if file_has_data: loaded_yuzde = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

            if loaded_kalite: self.guncelle_no_sutunu(self.tree_kalite, True)
            if loaded_yuzde: self.guncelle_no_sutunu(self.tree_yuzde, False)
        finally:
            self.statusbar_guncelle()

    def manuel_yedek_yukle(self): # Yedek dosyalarını seç ve manuel olarak yükle
        if not messagebox.askokcancel("Onay", "Mevcut tablo verileri silinecek ve seçilen yedekler yüklenecektir.\nDevam etmek istiyor musunuz?", parent=self.master):
            return

        self.master.update_idletasks()
        kalite_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Kalite_Kontrol_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]
        yuzde_files = [f for f in glob.glob(os.path.join(BACKUP_DIR, "Yuzde_Sapma_Olcumleri_Yedek_*.csv")) if not os.path.basename(f).startswith('~$') and os.path.getsize(f) > 0]

        if not kalite_files and not yuzde_files:
            messagebox.showinfo("Bilgi", "Yüklenecek (boş olmayan) yedek dosya bulunamadı.", parent=self.master)
            return

        result = self.yedek_sec_ve_yukle_dialog(kalite_files, yuzde_files)
        if result:
            load_kalite_files, load_yuzde_files = result
            if not load_kalite_files and not load_yuzde_files: return

            loaded_kalite, loaded_yuzde = False, False
            for item in self.tree_kalite.get_children(): self.tree_kalite.delete(item)
            for item in self.tree_yuzde.get_children(): self.tree_yuzde.delete(item)
            self.measurement_no_kalite, self.measurement_no_yuzde = 1, 1

            expected_kalite_headers = self.tree_kalite['columns']
            for file_path in load_kalite_files:
                try:
                    if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                        with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                            reader = csv.reader(f, delimiter=';'); headers = next(reader, None)
                            if headers and tuple(headers) == expected_kalite_headers:
                                file_has_data = False
                                for row in reader:
                                    if len(row) == len(expected_kalite_headers): self.tree_kalite.insert("", "end", values=row); file_has_data = True
                                if file_has_data: loaded_kalite = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)

            expected_yuzde_headers = self.tree_yuzde['columns']
            sapma_str_idx = list(expected_yuzde_headers).index("% Sapma Oranı") if "% Sapma Oranı" in expected_yuzde_headers else -1
            for file_path in load_yuzde_files:
                try:
                    if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                        with open(file_path, 'r', newline='', encoding='utf-8-sig') as f:
                            reader = csv.reader(f, delimiter=';'); headers = next(reader, None)
                            if headers and tuple(headers) == expected_yuzde_headers:
                                file_has_data = False
                                for i, row in enumerate(reader, start=1):
                                    if len(row) == len(expected_yuzde_headers):
                                        tags_to_apply = ()
                                        if sapma_str_idx != -1:
                                            try:
                                                sapma_val = float(row[sapma_str_idx].replace('%','').strip())
                                                if sapma_val > 9.99: tags_to_apply = ('high_deviation_tree',)
                                            except: pass
                                        self.tree_yuzde.insert("", "end", values=row, tags=tags_to_apply)
                                        file_has_data = True
                                if file_has_data: loaded_yuzde = True
                except Exception as e: messagebox.showerror("CSV Yükleme Hatası", f"'{os.path.basename(file_path)}' yüklenirken hata: {e}", parent=self.master)                
            if loaded_kalite: self.guncelle_no_sutunu(self.tree_kalite, True)
            if loaded_yuzde: self.guncelle_no_sutunu(self.tree_yuzde, False)
            self.statusbar_guncelle()


    def yedek_sec_ve_yukle_dialog(self, kalite_files, yuzde_files):
        dialog = tk.Toplevel(self.master)
        dialog.title("Yedek Dosyalarını Seçin")
        dialog.geometry("600x400")
        dialog.transient(self.master)
        dialog.grab_set()
        dialog.resizable(False, False)

        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(expand=True, fill="both")

        ttk.Label(main_frame, text="Yüklemek istediğiniz yedek dosyalarını seçin:", justify=tk.LEFT).pack(pady=(0,10), anchor='w')

        kalite_frame = ttk.LabelFrame(main_frame, text="Kalite Kontrol Yedekleri", padding="10")
        kalite_frame.pack(pady=5, fill="x")
        kalite_inner_frame = ttk.Frame(kalite_frame)
        kalite_inner_frame.pack(fill="both", expand=True)
        kalite_scrollbar = ttk.Scrollbar(kalite_inner_frame, orient=tk.VERTICAL)
        kalite_listbox = tk.Listbox(kalite_inner_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, yscrollcommand=kalite_scrollbar.set)
        kalite_scrollbar.config(command=kalite_listbox.yview)
        kalite_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        kalite_listbox.pack(side=tk.LEFT, fill="both", expand=True)
        kalite_map = {os.path.basename(f): f for f in kalite_files}
        for f_basename in sorted(kalite_map.keys(), reverse=True):
            kalite_listbox.insert(tk.END, f_basename)

        yuzde_frame = ttk.LabelFrame(main_frame, text="Yüzde Sapma Yedekleri", padding="10")
        yuzde_frame.pack(pady=5, fill="x")
        yuzde_inner_frame = ttk.Frame(yuzde_frame)
        yuzde_inner_frame.pack(fill="both", expand=True)
        yuzde_scrollbar = ttk.Scrollbar(yuzde_inner_frame, orient=tk.VERTICAL)
        yuzde_listbox = tk.Listbox(yuzde_inner_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, yscrollcommand=yuzde_scrollbar.set)
        yuzde_scrollbar.config(command=yuzde_listbox.yview)
        yuzde_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        yuzde_listbox.pack(side=tk.LEFT, fill="both", expand=True)
        yuzde_map = {os.path.basename(f): f for f in yuzde_files}
        for f_basename in sorted(yuzde_map.keys(), reverse=True):
            yuzde_listbox.insert(tk.END, f_basename)

        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=15)

        result = {"kalite": None, "yuzde": None}

        def on_ok():
            selected_kalite_indices = kalite_listbox.curselection()
            selected_yuzde_indices = yuzde_listbox.curselection()
            result["kalite"] = [kalite_map[kalite_listbox.get(i)] for i in selected_kalite_indices]
            result["yuzde"] = [yuzde_map[yuzde_listbox.get(i)] for i in selected_yuzde_indices]
            dialog.destroy()

        def on_cancel():
            result["kalite"] = None
            result["yuzde"] = None
            dialog.destroy()

        ok_button = ttk.Button(button_frame, text="Seçilenleri Yükle", command = on_ok)
        ok_button.pack(side=tk.LEFT, padx=10)
        cancel_button = ttk.Button(button_frame, text="İptal Et / Boş Başlat", command=on_cancel)
        cancel_button.pack(side=tk.LEFT, padx=10)

        dialog.protocol("WM_DELETE_WINDOW", on_cancel)
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')

        dialog.wait_window()

        if result["kalite"] is not None and result["yuzde"] is not None:
            return result["kalite"], result["yuzde"]
        else:
             return None

    def birimleri_comboboxa_yukle(self):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("SELECT birim_adi FROM birimler")
        results = cursor.fetchall()
        self.birimler = [row[0] for row in results]
        conn.close()
        self.cmb_birim['values'] = self.birimler
        if self.birimler:
            self.cmb_birim.current(0)

    def cihaz_tiplerini_comboboxa_yukle(self): # Veritabanından Cihaz Tiplerini al ve Combobox a yükle
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("SELECT cihaz_tipi FROM cihaz_tipleri")
        results = cursor.fetchall()
        self.device_types = sorted([row[0] for row in results], key=locale.strxfrm)
        conn.close()
        self.cmb_device_type['values'] = self.device_types

        last_type = self.program_ayari_yukle("last_selected_device_type")
        if last_type and last_type in self.device_types:
            self.cmb_device_type.set(last_type)
        else:
            self.cmb_device_type.set("")

    def cihaz_seri_no_lar_comboboxa_yukle(self): # Veritabanından Seçili cihaz tipine ait Cihaz Seri Numaraları al ve Combobox a yükle
        # Seçili cihaz tipini al
        selected_type = self.cmb_device_type.get()
        
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        try:
            if selected_type:
                cursor.execute("""
                    SELECT cihaz_seri 
                    FROM cihaz_serileri 
                    WHERE cihaz_tipi = ?
                    """, (selected_type,))
            else:
                cursor.execute("SELECT cihaz_seri FROM cihaz_serileri") # Hiçbir cihaz tipi seçili değilse tüm seri numaralarını getir
                
            results = cursor.fetchall()
            self.device_serials = sorted([row[0] for row in results], key=locale.strxfrm)
            self.cmb_device_serial['values'] = self.device_serials
            
            # Combobox'ı güncelle
            if self.device_serials:
                self.cmb_device_serial.current(0)
            else:
                self.cmb_device_serial.set("")
                
        except sqlite3.Error as e:
            messagebox.showerror("Veritabanı Hatası", f"Seri numaraları yüklenirken hata: {e}")
        finally:
            conn.close()

        default_serial = "BG709223125" # Varsayılan olarak bu Seri Numarasını getir
        if default_serial in self.device_serials:
            self.cmb_device_serial.set(default_serial)
        elif self.device_serials:
            self.cmb_device_serial.current(0) # Eğer "BG709223125" yoksa ilk değeri seç
        else:
            self.cmb_device_serial.set("") # Eğer hiç veri yoksa boş bırak

    def radyo_istasyonlar_comboboxa_yukle(self):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("SELECT radyo_adi, radyo_url FROM radyolar")
        results = cursor.fetchall()
        local_radio_station_names = [row[0] for row in results]
        local_radio_station_map = {row[0]: row[1] for row in results}
        conn.close()

        self.radio_station_names = local_radio_station_names
        self.radio_station_map = local_radio_station_map
        return local_radio_station_names, local_radio_station_map

    def get_son4hane_for_device(self, birim, tip, seri):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT son_4_hane FROM cihaz_kayitlari
            WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ?
            ORDER BY id DESC LIMIT 1
        """, (birim, tip, seri))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else ""

    def get_son4hane_list_for_device(self, birim, tip, seri):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT DISTINCT son_4_hane FROM cihaz_kayitlari
            WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ?
        """, (birim, tip, seri))
        results = cursor.fetchall()
        conn.close()
        return [row[0] for row in results] if results else []

    def birim_ekle_pencere(self):
        def tamam_click():
            yeni_birim = entry.get().strip().upper()
            if not yeni_birim:
                messagebox.showwarning("Uyarı", "Birim adı boş olamaz!", parent=top)
                return
            if not all(c.isalpha() or c.isspace() or c in "ĞÜŞİÖÇ" for c in yeni_birim if c.strip()):
                messagebox.showwarning("Uyarı", "Birim adı sadece harf, Türkçe karakterler ve boşluk içerebilir!", parent=top)
                return
            if yeni_birim in self.birimler:
                messagebox.showwarning("Uyarı", "Birim zaten mevcut!", parent=top)
                return

            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                cursor.execute("INSERT INTO birimler (birim_adi) VALUES (?)", (yeni_birim,))
                conn.commit()
                conn.close()
                self.birimleri_comboboxa_yukle()
                self.cmb_birim.set(yeni_birim)
                top.destroy()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Birim eklenirken hata: {e}", parent=top)

        def on_key_release(event):
            current = entry.get()
            entry.delete(0, tk.END)
            entry.insert(0, current.upper())

        top = tk.Toplevel(self.master)
        top.title("Birim Ekle")
        top.geometry("350x130")
        top.resizable(False, False)
        top.transient(self.master)
        top.grab_set()

        if self.plus_icon:
            top.iconbitmap(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Resources", "plus.ico"))

        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f"{width}x{height}+{x}+{y}")

        ttk.Label(top, text="Yeni Birim Adı:").pack(pady=10)
        entry = ttk.Entry(top, width=40)
        entry.pack(pady=5)
        entry.bind("<KeyRelease>", on_key_release)
        ttk.Button(top, text="Tamam", command=tamam_click).pack(pady=15)
        entry.focus_set()
        top.wait_window()

    def birim_sil(self):
        secilen = self.cmb_birim.get()
        if not secilen:
            messagebox.showwarning("Uyarı", "Silinecek birim seçiniz!")
            return
        if messagebox.askokcancel("Onay", f"'{secilen}' birimini silmek istediğinize emin misiniz?"):
            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                cursor.execute("DELETE FROM birimler WHERE birim_adi = ?", (secilen,))
                conn.commit()
                conn.close()
                self.birimleri_comboboxa_yukle()
                if self.birimler:
                    self.cmb_birim.current(0)
                else:
                    self.cmb_birim.set("")
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Birim silinirken hata: {e}")

    def add_or_update_device_assignment(self, birim, tip, seri, son4):
        if not (birim and tip and seri and son4 and len(son4)==4):
            messagebox.showerror("Hata", "Lütfen tüm alanları doldurun ve Seri numarasının son 4 hanesini kontrol edin.")
            return False

        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        try:
            # Aynı birim, seri, son4hane başka bir cihaz tipine kayıtlı mı?
            cursor.execute("""
                SELECT cihaz_tipi FROM cihaz_kayitlari
                WHERE birim_adi = ? AND cihaz_seri = ? AND son_4_hane = ? AND cihaz_tipi != ?
            """, (birim, seri, son4, tip))
            existing_assignment = cursor.fetchone()
            if existing_assignment:
                messagebox.showerror(
                    "Kayıt Hatası",
                    f" '{existing_assignment[0]}' cihazı, '{seri}{son4}' seri numarasıyla zaten '{birim}' biriminde kayıtlı.\n"
                    "Cihaz Tipi, Seri No ve Son 4 Haneyi verilerini tekrar kontrol edin."
                )
                return False

            # Eski kontrol: Aynı cihaz başka bir birime kayıtlı mı?
            cursor.execute("""
                SELECT birim_adi FROM cihaz_kayitlari
                WHERE cihaz_tipi = ? AND cihaz_seri = ? AND son_4_hane = ? AND birim_adi != ?
            """, (tip, seri, son4, birim))
            existing_assignment_birim = cursor.fetchone()
            if existing_assignment_birim:
                messagebox.showerror(
                    "Kayıt Hatası",
                    f"Bu Seri Numaralı cihaz ({seri}{son4}) zaten '{existing_assignment_birim[0]}' birimine kayıtlı.\n"
                    "Girdiğiniz Son 4 Haneyi tekrar kontrol ediniz."
                )
                return False

            cursor.execute("""
                INSERT OR IGNORE INTO cihaz_kayitlari (birim_adi, cihaz_tipi, cihaz_seri, son_4_hane)
                VALUES (?, ?, ?, ?)
            """, (birim, tip, seri, son4))
            conn.commit()
            self.on_birim_cihaz_secildi()
            return True
        except sqlite3.Error as e:
            messagebox.showerror("Veritabanı Hatası", f"Cihaz kaydı sırasında hata: {e}")
            return False
        finally:
            conn.close()

    def kontrol_cihaz_kayit_cakisma(self, event=None):
        if getattr(self, "cakisma_uyarildi", False):
            return  # Zaten uyarı gösterildi, tekrar gösterme

        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        son4 = self.cmb_son4hane.get().strip().upper()
        if not (birim and tip and seri and son4 and len(son4) == 4):
            return  # Tüm alanlar dolu değilse kontrol etme

        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT birim_adi, cihaz_tipi FROM cihaz_kayitlari
                WHERE cihaz_seri = ? AND son_4_hane = ? AND (birim_adi != ? OR cihaz_tipi != ?)
            """, (seri, son4, birim, tip))
            result = cursor.fetchone()
            if result:
                self.cakisma_uyarildi = True
                messagebox.showerror(
                    "Kayıt Çakışması",
                    f"{seri}{son4} seri numarasına sahip {result[1]} cihazı zaten {result[0]} biriminde kayıtlı.\n"
                    "Girdiğiniz verileri tekrar kontrol ediniz."
                )
                self.master.after(100, lambda: setattr(self, "cakisma_uyarildi", False))
                self.master.focus()
                return "break"
        finally:
            conn.close()

    def check_device_availability(self, birim, tip, seri, son4):
        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT birim_adi FROM cihaz_kayitlari
            WHERE cihaz_tipi = ? AND cihaz_seri = ? AND son_4_hane = ? AND birim_adi != ?
        """, (tip, seri, son4, birim))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else None

    def on_birim_cihaz_secildi(self, event=None):
        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        if birim and tip and seri:
            son4_list = self.get_son4hane_list_for_device(birim, tip, seri)
            self.cmb_son4hane['values'] = son4_list
            if son4_list:
                self.cmb_son4hane.set(son4_list[0])
            else:
                self.cmb_son4hane.set("")
            self.validate_son4hane_input()
        else:
            self.cmb_son4hane['values'] = []
            self.cmb_son4hane.set("")
            self.validate_son4hane_input()

    def on_son4hane_changed(self, event=None):
        self.validate_son4hane_input()

    def on_seri_no_entered(self, event=None):
        new_serial = self.cmb_device_serial.get().strip().upper()
        if new_serial and new_serial not in self.device_serials:
            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                cursor.execute("INSERT OR IGNORE INTO cihaz_serileri (cihaz_seri) VALUES (?)", (new_serial,))
                conn.commit()
                conn.close()
                self.cihaz_seri_no_lar_comboboxa_yukle()
                self.cmb_device_serial.set(new_serial)
                self.on_birim_cihaz_secildi()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz seri no eklenirken hata: {e}")

    def on_device_type_selected(self, event=None):
        selected_type = self.cmb_device_type.get()
        self.program_ayari_kaydet("last_selected_device_type", selected_type)
        self.cihaz_seri_no_lar_comboboxa_yukle() # Cihaz tipi değiştiğinde seri numaralarını güncelle

    def on_device_serial_selected(self, event=None):
        selected_serial = self.cmb_device_serial.get()
        self.program_ayari_kaydet("last_selected_serial_no", selected_serial)

    def validate_son4hane_input(self, event=None):
        content = self.cmb_son4hane.get().upper()
        new_content = "".join(filter(str.isalnum, content))[:4]
        if self.cmb_son4hane.get() != new_content:
            self.cmb_son4hane.set(new_content)

        if len(new_content) == 4:
            self.cmb_son4hane.config(style="Valid.TCombobox")
            if hasattr(self, 'son4hane_tooltip') and self.son4hane_tooltip.tooltip_window:
                self.son4hane_tooltip.hide_tooltip()
        elif len(new_content) > 0:
            self.cmb_son4hane.config(style="Invalid.TCombobox")
            if not hasattr(self, 'son4hane_tooltip') or self.son4hane_tooltip.widget != self.cmb_son4hane:
                self.son4hane_tooltip = ToolTip(self.cmb_son4hane, "Cihazın seri numarasının son 4 hanesini giriniz ya da varsa listeden seçiniz.")
        else:
            self.cmb_son4hane.config(style="TCombobox")
            if hasattr(self, 'son4hane_tooltip') and self.son4hane_tooltip.tooltip_window:
                self.son4hane_tooltip.hide_tooltip()
        return True

    def validate_l_entry(self, P, widget_name, min_val, max_val):
        widget = self.master.nametowidget(widget_name)
        if P == "":
            widget.config(style="TEntry")
            if widget_name in self.l_entry_tooltips and self.l_entry_tooltips[widget_name].tooltip_window:
                self.l_entry_tooltips[widget_name].hide_tooltip()
            return True
        try:
            value = int(P)
            if min_val <= value <= max_val:
                widget.config(style="Valid.TEntry")
                if widget_name in self.l_entry_tooltips and self.l_entry_tooltips[widget_name].tooltip_window:
                    self.l_entry_tooltips[widget_name].hide_tooltip()
                return True
            else:
                widget.config(style="Invalid.TEntry")
                if widget_name in self.l_entry_tooltips:
                    self.l_entry_tooltips[widget_name].show_tooltip()
                return True
        except ValueError:
            widget.config(style="Invalid.TEntry")
            if widget_name in self.l_entry_tooltips:
                self.l_entry_tooltips[widget_name].show_tooltip()
            return True

    def check_l_entries_valid_for_aktar(self):
        valid_l1, valid_l2, valid_l3 = True, True, True
        try:
            l1_val = int(self.txt_l1.get())
            if not (36 <= l1_val <= 108): valid_l1 = False
        except ValueError: valid_l1 = False
        if not self.txt_l1.get(): valid_l1 = False

        try:
            l2_val = int(self.txt_l2.get())
            if not (144 <= l2_val <= 216): valid_l2 = False
        except ValueError: valid_l2 = False
        if not self.txt_l2.get(): valid_l2 = False

        try:
            l3_val = int(self.txt_l3.get())
            if not (252 <= l3_val <= 396): valid_l3 = False
        except ValueError: valid_l3 = False
        if not self.txt_l3.get(): valid_l3 = False

        if not (valid_l1 and valid_l2 and valid_l3):
            messagebox.showerror("Hata", "L1, L2, L3 değerlerini kontrol edin ve referans aralıklarında değerler girin.")
            return None
        return l1_val, l2_val, l3_val

    def cihaz_markasi_ekle_pencere(self):
        def tamam_click():
            yeni_cihaz = entry.get().strip().upper()
            if not yeni_cihaz:
                messagebox.showwarning("Uyarı", "Cihaz adı boş olamaz!", parent=top)
                return
            
            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                cursor.execute("INSERT INTO cihaz_tipleri (cihaz_tipi) VALUES (?)", (yeni_cihaz,))
                conn.commit()
                conn.close()
                self.cihaz_tiplerini_comboboxa_yukle()
                self.cmb_device_type.set(yeni_cihaz)
                top.destroy()
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz markası eklenirken hata: {e}", parent=top)

        def on_key_release(event):
            current = entry.get()
            entry.delete(0, tk.END)
            entry.insert(0, current.upper())

        top = tk.Toplevel(self.master)
        top.title("Cihaz Markası Ekle")
        top.geometry("350x130")
        top.resizable(False, False)
        top.transient(self.master)
        top.grab_set()

        if self.plus_icon:
            top.iconbitmap(os.path.join(os.path.dirname(os.path.abspath(__file__)), "Resources", "plus.ico"))

        ttk.Label(top, text="Yeni Cihaz Markası:").pack(pady=10)
        entry = ttk.Entry(top, width=40)
        entry.insert(0, "GLUKOMETRE-")
        entry.pack(pady=5)
        entry.bind("<KeyRelease>", on_key_release)
        ttk.Button(top, text="Tamam", command=tamam_click).pack(pady=15)
        
        entry.focus_set()
        entry.icursor(tk.END)  # İmleci en sona konumlandır
        
        # Pencereyi ortala
        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f"{width}x{height}+{x}+{y}")
        top.wait_window()

    def cihaz_markasi_sil(self):
        secilen = self.cmb_device_type.get()
        if not secilen:
            messagebox.showwarning("Uyarı", "Silinecek cihaz markasını seçiniz!")
            return
                
        if messagebox.askokcancel("Onay", f"'{secilen}' cihaz markasını silmek istediğinize emin misiniz?"):
            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                cursor.execute("DELETE FROM cihaz_tipleri WHERE cihaz_tipi = ?", (secilen,))
                conn.commit()
                conn.close()
                self.cihaz_tiplerini_comboboxa_yukle()
                if self.device_types:
                    self.cmb_device_type.current(0)
                else:
                    self.cmb_device_type.set("")
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz markası silinirken hata: {e}")

    def open_cihaz_arama_dialog(self):
        dialog = tk.Toplevel(self.master)
        dialog.title("Cihaz Arama")
        dialog.geometry("800x550")
        dialog.transient(self.master)
        dialog.grab_set()

        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            search_icon_path = os.path.join(script_dir, "Resources", "search.ico")
            dialog.iconbitmap(search_icon_path)
        except:
            pass

        # Ana Frame
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill="both", expand=True)

        # Arama Kriterleri Frame
        search_frame = ttk.LabelFrame(main_frame, text="Arama Kriterleri", padding="10")
        search_frame.pack(fill="x", padx=5, pady=(0,10))

        # Comboboxlar için Frame
        criteria_frame = ttk.Frame(search_frame)
        criteria_frame.pack(fill="x", expand=True)

        # Grid düzeni için column ağırlıkları
        criteria_frame.columnconfigure(1, weight=1)
        criteria_frame.columnconfigure(3, weight=1)

        # Birim Seçimi
        ttk.Label(criteria_frame, text="Birim:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        cmb_birim_search = ttk.Combobox(criteria_frame, state="normal")
        cmb_birim_search['values'] = self.birimler
        ToolTip(cmb_birim_search, "Listeden bir Birim seçerek, o birimde kayıtlı tüm cihazları görebilirsiniz.")
        cmb_birim_search.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        
        # Cihaz Tipi Seçimi
        ttk.Label(criteria_frame, text="Cihaz Tipi:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        cmb_cihaz_tipi_search = ttk.Combobox(criteria_frame, state="normal")
        cmb_cihaz_tipi_search['values'] = self.device_types
        ToolTip(cmb_cihaz_tipi_search, "Listeden bir Cihaz Tipi - Markası seçerek, o cihaz markasının hangi birimlerde bulunduğunu görebilirsiniz.")
        cmb_cihaz_tipi_search.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        # Seri No Seçimi
        ttk.Label(criteria_frame, text="Seri No:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        cmb_seri_no_search = ttk.Combobox(criteria_frame, state="normal")
        ToolTip(cmb_seri_no_search, "Listeden bir Seri numarası seçerek ve son 4 hane girerek, o cihazın hangi birime ait olduğunu bulabilirsiniz.")
        cmb_seri_no_search.grid(row=2, column=1, padx=5, pady=5, sticky="ew")
        
        # Son 4 Hane
        ttk.Label(criteria_frame, text="Son 4 Hane:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        cmb_son4hane_search = ttk.Combobox(criteria_frame, state="normal")
        cmb_son4hane_search.grid(row=3, column=1, padx=5, pady=5, sticky="ew")

        # Sonuçlar için Treeview
        columns = ("Birim", "Cihaz Tipi", "Seri No")  # Kayıt Tarihi sütununu kaldırdık
        tree_results = ttk.Treeview(main_frame, columns=columns, show="headings", selectmode="browse")
        tree_results.pack(fill="both", expand=True, pady=(0,10))

        # Scrollbarlar
        vsb = ttk.Scrollbar(main_frame, orient="vertical", command=tree_results.yview)
        vsb.pack(side="right", fill="y")
        hsb = ttk.Scrollbar(main_frame, orient="horizontal", command=tree_results.xview)
        hsb.pack(side="bottom", fill="x")
        tree_results.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        # Tablo başlıkları ve genişlikleri - Tüm sütunları ortala
        widths = [190, 170, 150]
        for col, width in zip(columns, widths):
            tree_results.heading(col, text=col, 
                            command=lambda c=col: self.treeview_sort_column(tree_results, c, False))
            tree_results.column(col, width=width, anchor=tk.CENTER, minwidth=width)  # anchor=tk.CENTER ile ortala

        def update_results(search_type=None):
            for item in tree_results.get_children():
                tree_results.delete(item)

            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()

            try:
                birim = cmb_birim_search.get().strip()
                cihaz_tipi = cmb_cihaz_tipi_search.get().strip()
                seri_no = cmb_seri_no_search.get().strip()
                son4hane = cmb_son4hane_search.get().strip()

                if birim and cihaz_tipi:
                    # Birim ve Cihaz Tipi ile arama
                    cursor.execute("""
                        SELECT birim_adi, cihaz_tipi, cihaz_seri, son_4_hane
                        FROM cihaz_kayitlari 
                        WHERE birim_adi = ? AND cihaz_tipi = ?
                        ORDER BY birim_adi, cihaz_tipi, cihaz_seri
                    """, (birim, cihaz_tipi))
                elif birim:
                    cursor.execute("""
                        SELECT birim_adi, cihaz_tipi, cihaz_seri, son_4_hane
                        FROM cihaz_kayitlari 
                        WHERE birim_adi = ?
                        ORDER BY cihaz_tipi, cihaz_seri
                    """, (birim,))
                elif cihaz_tipi:
                    cursor.execute("""
                        SELECT birim_adi, cihaz_tipi, cihaz_seri, son_4_hane
                        FROM cihaz_kayitlari 
                        WHERE cihaz_tipi = ?
                        ORDER BY birim_adi, cihaz_seri
                    """, (cihaz_tipi,))
                elif seri_no and son4hane:
                    cursor.execute("""
                        SELECT birim_adi, cihaz_tipi, cihaz_seri, son_4_hane
                        FROM cihaz_kayitlari 
                        WHERE cihaz_seri = ? AND son_4_hane = ?
                        ORDER BY birim_adi, cihaz_tipi
                    """, (seri_no, son4hane))

                results = cursor.fetchall()
                for row in results:
                    birim_adi, cihaz_tipi, cihaz_seri, son4hane = row
                    tam_seri_no = f"{cihaz_seri}{son4hane}"
                    tree_results.insert("", "end", values=(birim_adi, cihaz_tipi, tam_seri_no))

            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Arama yapılırken hata oluştu: {e}")
            finally:
                conn.close()

        def on_birim_selected(event=None):
            update_results("birim")
        def on_cihaz_tipi_selected(event=None):
            update_results("cihaz_tipi")
        def on_seri_no_selected(event=None):
            update_results("seri_no")
        def on_son4hane_selected(event=None):
            update_results("son4hane")

        cmb_birim_search.bind('<<ComboboxSelected>>', on_birim_selected)
        cmb_cihaz_tipi_search.bind('<<ComboboxSelected>>', on_cihaz_tipi_selected)
        cmb_seri_no_search.bind('<<ComboboxSelected>>', on_seri_no_selected)
        cmb_son4hane_search.bind('<<ComboboxSelected>>', on_son4hane_selected)

        cmb_birim_search.bind('<KeyRelease>', on_birim_selected)
        cmb_cihaz_tipi_search.bind('<KeyRelease>', on_cihaz_tipi_selected)
        cmb_seri_no_search.bind('<KeyRelease>', on_seri_no_selected)
        cmb_son4hane_search.bind('<KeyRelease>', on_son4hane_selected)

        # Pencereyi ortala
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        dialog.focus_set()
        dialog.wait_window()

    def open_cihaz_ekle_sil_dialog(self):
        dialog = tk.Toplevel(self.master)
        dialog.title("Cihaz Ekle / Sil")
        dialog.geometry("800x310")
        dialog.transient(self.master)
        dialog.grab_set()

        # İkonu ayarla
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            plus_icon_path = os.path.join(script_dir, "Resources", "plus.ico")
            dialog.iconbitmap(plus_icon_path)
        except:
            pass

        # Ana Frame
        main_frame = ttk.Frame(dialog)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)

        # Cihaz Ekle Frame
        ekle_frame = ttk.LabelFrame(main_frame, text="Envantere Cihaz Ekle", padding=10)
        ekle_frame.pack(side="left", fill="both", expand=True, padx=(0, 5))
        ekle_frame.columnconfigure(1, weight=1)

        # Cihaz Sil Frame
        sil_frame = ttk.LabelFrame(main_frame, text="Envanterden Cihaz Sil", padding=10)
        sil_frame.pack(side="left", fill="both", expand=True, padx=(5, 0))
        sil_frame.columnconfigure(1, weight=1)

        # 1 - Birim Combobox (Sil)
        ttk.Label(sil_frame, text="Birim:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.cmb_birim_sil = ttk.Combobox(sil_frame, state="readonly")
        self.cmb_birim_sil['values'] = self.birimler  # Veritabanından birimler
        if self.birimler:
            self.cmb_birim_sil.current(0)
        self.cmb_birim_sil.grid(row=0, column=1, sticky="ew", padx=5, pady=5)

        # 2 - Cihaz Combobox (Sil)
        ttk.Label(sil_frame, text="Cihaz Tipi - Marka:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.cmb_cihaz_sil = ttk.Combobox(sil_frame, state="readonly")
        self.cmb_cihaz_sil.grid(row=1, column=1, sticky="ew", padx=5, pady=5)

        # 3 - Seri No Combobox (Sil)
        ttk.Label(sil_frame, text="Seri No:").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.cmb_seri_no_sil = ttk.Combobox(sil_frame, state="readonly")
        self.cmb_seri_no_sil.grid(row=2, column=1, sticky="ew", padx=5, pady=5)

        # 4 - Son 4 Hane Combobox (Sil)
        ttk.Label(sil_frame, text="Son 4 Hane:").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.cmb_son4hane_sil = ttk.Combobox(sil_frame, state="readonly")
        self.cmb_son4hane_sil.grid(row=3, column=1, sticky="ew", padx=5, pady=5)

        def load_cihazlar_sil(event=None):
            birim = self.cmb_birim_sil.get()
            if not birim:
                self.cmb_cihaz_sil['values'] = []
                self.cmb_cihaz_sil.set("")
                return

            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT cihaz_tipi FROM cihaz_kayitlari
                WHERE birim_adi = ?
            """, (birim,))
            cihazlar = [row[0] for row in cursor.fetchall()]
            conn.close()
            self.cmb_cihaz_sil['values'] = cihazlar
            if cihazlar:
                self.cmb_cihaz_sil.current(0)
                load_seri_no_sil()
            else:
                self.cmb_cihaz_sil.set("")
                self.cmb_seri_no_sil['values'] = []
                self.cmb_seri_no_sil.set("")
                self.cmb_son4hane_sil['values'] = []
                self.cmb_son4hane_sil.set("")

        def load_seri_no_sil(event=None):
            birim = self.cmb_birim_sil.get()
            cihaz = self.cmb_cihaz_sil.get()
            if not (birim and cihaz):
                self.cmb_seri_no_sil['values'] = []
                self.cmb_seri_no_sil.set("")
                return

            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT cihaz_seri FROM cihaz_kayitlari
                WHERE birim_adi = ? AND cihaz_tipi = ?
            """, (birim, cihaz))
            seri_nolar = [row[0] for row in cursor.fetchall()]
            conn.close()
            self.cmb_seri_no_sil['values'] = seri_nolar
            if seri_nolar:
                self.cmb_seri_no_sil.current(0)
                load_son4hane_sil()
            else:
                self.cmb_seri_no_sil.set("")
                self.cmb_son4hane_sil['values'] = []
                self.cmb_son4hane_sil.set("")

        def load_son4hane_sil(event=None):
            birim = self.cmb_birim_sil.get()
            cihaz = self.cmb_cihaz_sil.get()
            seri_no = self.cmb_seri_no_sil.get()
            if not (birim and cihaz and seri_no):
                self.cmb_son4hane_sil['values'] = []
                self.cmb_son4hane_sil.set("")
                return

            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT son_4_hane FROM cihaz_kayitlari
                WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ?
            """, (birim, cihaz, seri_no))
            son4haneler = [row[0] for row in cursor.fetchall()]
            conn.close()
            self.cmb_son4hane_sil['values'] = son4haneler
            if son4haneler:
                self.cmb_son4hane_sil.current(0)
            else:
                self.cmb_son4hane_sil.set("")

        # 5 - Cihaz Sil Butonu
        def cihaz_sil_click():
            birim = self.cmb_birim_sil.get()
            cihaz = self.cmb_cihaz_sil.get()
            seri_no = self.cmb_seri_no_sil.get()
            son4hane = self.cmb_son4hane_sil.get()
            cihaz_seri_no_tam = seri_no + son4hane

            if not (birim and cihaz and seri_no and son4hane):
                messagebox.showerror("Hata", "Lütfen tüm alanları seçin!")
                return

            if messagebox.askokcancel("Onay", f"{birim} birimine ait {cihaz_seri_no_tam} seri numaralı {cihaz} cihazını silmek istediğinize emin misiniz?"):
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                        DELETE FROM cihaz_kayitlari
                        WHERE birim_adi = ? AND cihaz_tipi = ? AND cihaz_seri = ? AND son_4_hane = ?
                    """, (birim, cihaz, seri_no, son4hane))
                    conn.commit()
                    messagebox.showinfo("Başarılı", "Cihaz başarıyla silindi.")
                    self.cihaz_tiplerini_comboboxa_yukle()
                    self.cihaz_seri_no_lar_comboboxa_yukle()
                    self.on_birim_cihaz_secildi()
                    self.statusbar_guncelle()
                    dialog.destroy()
                except sqlite3.Error as e:
                    messagebox.showerror("Veritabanı Hatası", f"Cihaz silinirken hata: {e}")
                finally:
                    conn.close()

        ttk.Button(sil_frame, text="Cihaz Sil", command=cihaz_sil_click).grid(row=4, column=1, sticky="e", padx=5, pady=10)

        self.cmb_birim_sil.bind("<<ComboboxSelected>>", load_cihazlar_sil)
        self.cmb_cihaz_sil.bind("<<ComboboxSelected>>", load_seri_no_sil)
        self.cmb_seri_no_sil.bind("<<ComboboxSelected>>", load_son4hane_sil)

        # 1 - Birim Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Birim:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.cmb_birim_ekle = ttk.Combobox(ekle_frame, state="readonly")
        self.cmb_birim_ekle['values'] = self.birimler  # Ana penceredeki birimler
        if self.birimler:
            self.cmb_birim_ekle.current(0)
        self.cmb_birim_ekle.grid(row=0, column=1, sticky="ew", padx=5, pady=5)

        # 2 - Cihaz Tipi - Marka Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Cihaz Tipi - Marka:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.cmb_cihaz_ekle = ttk.Combobox(ekle_frame, state="readonly")
        self.cmb_cihaz_ekle['values'] = self.device_types  # Ana penceredeki cihaz tipleri
        if self.device_types:
            self.cmb_cihaz_ekle.current(0)
        self.cmb_cihaz_ekle.grid(row=1, column=1, sticky="ew", padx=5, pady=5)
        self.cmb_cihaz_ekle.bind("<<ComboboxSelected>>", self.load_seri_no_ekle)

        # 3 - Seri No Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Seri No:").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.cmb_seri_no_ekle = ttk.Combobox(ekle_frame, state="normal")
        self.cmb_seri_no_ekle.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        # self.load_seri_no_ekle() # İlk cihaz tipine göre seri noları yükle

        # 4 - Son 4 Hane Textbox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Son 4 Hane:").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.txt_son4hane_ekle = ttk.Entry(ekle_frame)
        self.txt_son4hane_ekle.grid(row=3, column=1, sticky="ew", padx=5, pady=5)
        # self.txt_son4hane_ekle için KeyRelease event'i ve validate komutu orijinaldeki gibi kalmalı

        # 5 - Üretim Tarihi Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Üretim Tarihi:").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.cmb_uretim_tarihi = ttk.Combobox(ekle_frame, state="normal")
        self.cmb_uretim_tarihi.grid(row=4, column=1, sticky="ew", padx=5, pady=5)

        # 6 - Satın Alma Tarihi Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Satın Alma Tarihi:").grid(row=5, column=0, sticky="w", padx=5, pady=5)
        self.cmb_satin_alma_tarihi = ttk.Combobox(ekle_frame, state="normal")
        self.cmb_satin_alma_tarihi.grid(row=5, column=1, sticky="ew", padx=5, pady=5)

        # 7 - Temsilci Firma Combobox (ekle_frame içinde)
        ttk.Label(ekle_frame, text="Temsilci Firma:").grid(row=6, column=0, sticky="w", padx=5, pady=5)
        self.cmb_temsilci_firma = ttk.Combobox(ekle_frame, state="normal")
        self.cmb_temsilci_firma.grid(row=6, column=1, sticky="ew", padx=5, pady=5)

        try:
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()

            # Üretim Tarihi için (uretim_tarihleri tablosundan uretim_yili sütunundan)
            cursor.execute("SELECT uretim_yili FROM uretim_tarihleri ORDER BY uretim_yili")
            uretim_yillari_list = [r[0] for r in cursor.fetchall()]
            self.cmb_uretim_tarihi['values'] = uretim_yillari_list
            if not uretim_yillari_list: # Eğer liste boşsa, örnek bir değer veya boş string ekleyebilirsiniz.
                self.cmb_uretim_tarihi['values'] = ["Örn: 2023"]

            # Satın Alma Tarihi için (satin_alma_tarihleri tablosundan satin_alma_tarihi sütunundan)
            cursor.execute("SELECT satin_alma_tarihi FROM satin_alma_tarihleri ORDER BY satin_alma_tarihi")
            satin_alma_list = [r[0] for r in cursor.fetchall()]
            self.cmb_satin_alma_tarihi['values'] = satin_alma_list
            if not satin_alma_list:
                self.cmb_satin_alma_tarihi['values'] = ["Örn: 01.01.2024"]

            # Temsilci Firma için (temsilci_firmalar tablosundan temsilci_firma_adi sütunundan)
            cursor.execute("SELECT temsilci_firma_adi FROM temsilci_firmalar ORDER BY temsilci_firma_adi")
            firmalar_list = [r[0] for r in cursor.fetchall()]
            self.cmb_temsilci_firma['values'] = firmalar_list
            if not firmalar_list:
                 self.cmb_temsilci_firma['values'] = ["Örn: Firma Adı Ltd. Şti."]

            conn.close()
        except sqlite3.Error as e:
            # Programın çökmemesi için tablo/sütun yoksa hata mesajı yerine konsola yazdırılabilir.
            # Kullanıcıya bu tabloları oluşturması gerektiği hatırlatılabilir.
            print(f"Comboboxlar için özel veri yükleme hatası (Yeni tablolar/sütunlar eksik olabilir): {e}")
            # Geçici olarak boş değerler veya örnekler atanabilir
            if not hasattr(self.cmb_uretim_tarihi, 'cget') or not self.cmb_uretim_tarihi['values']:
                 self.cmb_uretim_tarihi['values'] = ["Veritabanı tablosu eksik olabilir."]
            if not hasattr(self.cmb_satin_alma_tarihi, 'cget') or not self.cmb_satin_alma_tarihi['values']:
                 self.cmb_satin_alma_tarihi['values'] = ["Veritabanı tablosu eksik olabilir."]
            if not hasattr(self.cmb_temsilci_firma, 'cget') or not self.cmb_temsilci_firma['values']:
                 self.cmb_temsilci_firma['values'] = ["Veritabanı tablosu eksik olabilir."]

        except Exception as e:
            print(f"Combobox veri yükleme sırasında genel hata: {e}")

        def son4hane_change_ekle(event): # İsim çakışmaması için _ekle eklendi
            current_text = self.txt_son4hane_ekle.get()
            self.txt_son4hane_ekle.delete(0, tk.END)
            self.txt_son4hane_ekle.insert(0, current_text.upper())
        self.txt_son4hane_ekle.bind("<KeyRelease>", son4hane_change_ekle)

        # Son 4 Hane için validate ve KeyRelease eventleri (orijinaldeki gibi)
        def son4hane_validate_ekle(new_text): # İsim çakışmaması için _ekle eklendi
            return len(new_text) <= 4 and all(c.isalnum() for c in new_text)
        son4hane_vcmd_ekle = (dialog.register(son4hane_validate_ekle), '%P')
        self.txt_son4hane_ekle.config(validate="key", validatecommand=son4hane_vcmd_ekle)

        def son4hane_validate(new_text):
            return all(c.isalnum() for c in new_text)

        son4hane_vcmd = (dialog.register(son4hane_validate), '%P')
        self.txt_son4hane_ekle.config(validate="key", validatecommand=son4hane_vcmd)

        def son4hane_change(event):
            current_text = self.txt_son4hane_ekle.get()
            self.txt_son4hane_ekle.delete(0, tk.END)
            self.txt_son4hane_ekle.insert(0, current_text.upper())

        self.txt_son4hane_ekle.bind("<KeyRelease>", son4hane_change)

        # 5 - Cihaz Ekle Butonu
        def cihaz_ekle_click():
            birim = self.cmb_birim_ekle.get()
            cihaz_tipi_marka = self.cmb_cihaz_ekle.get()
            seri_no_ana = self.cmb_seri_no_ekle.get().strip().upper()
            son_4_hane = self.txt_son4hane_ekle.get().strip().upper()

            uretim_yili_val = self.cmb_uretim_tarihi.get().strip()
            satin_alma_tarihi_val = self.cmb_satin_alma_tarihi.get().strip()
            temsilci_firma_adi_val = self.cmb_temsilci_firma.get().strip()
            
            # Gerekli alanların kontrolü
            if not (birim and cihaz_tipi_marka and seri_no_ana and son_4_hane):
                messagebox.showerror("Hata", "Birim, Cihaz Tipi, Seri No ve Son 4 Hane alanları doldurulmalıdır!", parent=dialog)
                return
            
            if len(son_4_hane) != 4:
                messagebox.showerror("Hata", "Son 4 Hane tam olarak 4 karakter olmalıdır!", parent=dialog)
                self.txt_son4hane_ekle.focus_set()
                return

            conn = None # bağlantıyı başta None olarak tanımla
            try:
                conn = sqlite3.connect(VERITABANI_DOSYASI)
                cursor = conn.cursor()

                # 1. `cihaz_serileri` tablosuna ekleme (cihaz_seri PRIMARY KEY, cihaz_tipi NOT NULL)
                # Bu tablo, bir seri numarasının hangi tip cihaza ait olduğunu genel olarak listeler.
                # Eğer aynı seri no farklı tipte eklenmeye çalışılırsa PK hatası verir, INSERT OR IGNORE bunu yakalar.
                cursor.execute("""
                    INSERT OR IGNORE INTO cihaz_serileri (cihaz_seri, cihaz_tipi) 
                    VALUES (?, ?)
                """, (seri_no_ana, cihaz_tipi_marka))

                # 2. `cihaz_kayitlari` tablosuna asıl cihaz envanter kaydını ekleme
                # Bu tablo UNIQUE constraint (cihaz_tipi, cihaz_seri, son_4_hane) ile aynı tam cihazın tekrar eklenmesini önler.
                sql_insert_cihaz_kayitlari = """
                    INSERT INTO cihaz_kayitlari 
                    (birim_adi, cihaz_tipi, cihaz_seri, son_4_hane, uretim_yili, satin_alma_tarihi, temsilci_firma)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                cursor.execute(sql_insert_cihaz_kayitlari, 
                               (birim, cihaz_tipi_marka, seri_no_ana, son_4_hane,
                                uretim_yili_val if uretim_yili_val and uretim_yili_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 2023"] else None,
                                satin_alma_tarihi_val if satin_alma_tarihi_val and satin_alma_tarihi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 01.01.2024"] else None,
                                temsilci_firma_adi_val if temsilci_firma_adi_val and temsilci_firma_adi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: Firma Adı Ltd. Şti."] else None))

                # 3. Yeni girilen Üretim Yılı, Satın Alma Tarihi, Temsilci Firma bilgilerini kendi tablolarına ekle (comboboxları zenginleştirmek için)
                if uretim_yili_val and uretim_yili_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 2023"]:
                    cursor.execute("INSERT OR IGNORE INTO uretim_tarihleri (uretim_yili) VALUES (?)", (uretim_yili_val,))
                
                if satin_alma_tarihi_val and satin_alma_tarihi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 01.01.2024"]:
                    cursor.execute("INSERT OR IGNORE INTO satin_alma_tarihleri (satin_alma_tarihi) VALUES (?)", (satin_alma_tarihi_val,))
                
                if temsilci_firma_adi_val and temsilci_firma_adi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: Firma Adı Ltd. Şti."]:
                    cursor.execute("INSERT OR IGNORE INTO temsilci_firmalar (temsilci_firma_adi) VALUES (?)", (temsilci_firma_adi_val,))
                
                conn.commit()
                messagebox.showinfo("Başarılı", f"{cihaz_tipi_marka} ({seri_no_ana}{son_4_hane}) cihazı\n'{birim}' birimine başarıyla eklendi.", parent=dialog)
                self.statusbar_guncelle()
                # Diyalogdaki combobox'ları güncelle (yeni eklenen değerler hemen görünsün diye)
                cursor.execute("SELECT uretim_yili FROM uretim_tarihleri ORDER BY uretim_yili")
                self.cmb_uretim_tarihi['values'] = [r[0] for r in cursor.fetchall()]
                
                cursor.execute("SELECT satin_alma_tarihi FROM satin_alma_tarihleri ORDER BY satin_alma_tarihi")
                self.cmb_satin_alma_tarihi['values'] = [r[0] for r in cursor.fetchall()]

                cursor.execute("SELECT temsilci_firma_adi FROM temsilci_firmalar ORDER BY temsilci_firma_adi")
                self.cmb_temsilci_firma['values'] = [r[0] for r in cursor.fetchall()]

                # İsteğe bağlı: Giriş alanlarını temizle ve son girilenleri seçili bırak
                self.cmb_seri_no_ekle.set("") 
                self.txt_son4hane_ekle.delete(0, tk.END)
                self.cmb_uretim_tarihi.set(uretim_yili_val if uretim_yili_val and uretim_yili_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 2023"] else "")
                self.cmb_satin_alma_tarihi.set(satin_alma_tarihi_val if satin_alma_tarihi_val and satin_alma_tarihi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: 01.01.2024"] else "")
                self.cmb_temsilci_firma.set(temsilci_firma_adi_val if temsilci_firma_adi_val and temsilci_firma_adi_val not in ["Veritabanı tablosu eksik olabilir.", "Örn: Firma Adı Ltd. Şti."] else "")
                
                self.cmb_seri_no_ekle.focus_set() # Yeni kayıt için seri no'ya odaklan

                # Ana penceredeki combobox'ları da güncellemek gerekebilir (programın genel yapısına göre)
                self.cihaz_tiplerini_comboboxa_yukle() 
                self.cihaz_seri_no_lar_comboboxa_yukle() 
                self.on_birim_cihaz_secildi()

            except sqlite3.IntegrityError as e:
                 messagebox.showerror("Veritabanı Hatası", f"Cihaz eklenirken bütünlük hatası: {e}\nBu cihaz (Cihaz Tipi+Seri No+Son 4 Hane kombinasyonu ile) zaten '{birim}' biriminde veya başka bir birimde kayıtlı olabilir.\nLütfen girdiğiniz bilgileri kontrol edin.", parent=dialog)
            except sqlite3.Error as e:
                messagebox.showerror("Veritabanı Hatası", f"Cihaz eklenirken hata: {e}", parent=dialog)
            finally:
                if conn:
                    conn.close()

        ttk.Button(ekle_frame, text="Cihaz Ekle", command=cihaz_ekle_click).grid(row=7, column=0, columnspan=2, sticky="e", padx=5, pady=10)

        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f"{width}x{height}+{x}+{y}")
        dialog.focus_set()
        dialog.wait_window()

    def load_seri_no_ekle(self, event=None):
        cihaz_tipi = self.cmb_cihaz_ekle.get()
        if not cihaz_tipi:
            self.cmb_seri_no_ekle['values'] = []
            return

        conn = sqlite3.connect(VERITABANI_DOSYASI)
        cursor = conn.cursor()
        cursor.execute("SELECT cihaz_seri FROM cihaz_serileri WHERE cihaz_tipi = ?", (cihaz_tipi,))
        seri_numaralari = [row[0] for row in cursor.fetchall()]
        conn.close()

        self.cmb_seri_no_ekle['values'] = seri_numaralari
        if seri_numaralari:
            self.cmb_seri_no_ekle.current(0)
        else:
            self.cmb_seri_no_ekle.set("")


    def cihaz_karsilastirma_formu_olustur(self):
        if not self.tree_yuzde.get_children():
            messagebox.showinfo("Veri Yok", "Yüzde Sapma tablosunda aktarılacak veri bulunmuyor.", parent=self.master)
            return

        SABLON_KARSILASTIRMA_DOSYASI = os.path.join(SABLONLAR_DIR, "GLUKOMETRE_CIHAZI_KARSILASTIRMA_SONUC_FORMU.xlsx")

        if not os.path.exists(SABLON_KARSILASTIRMA_DOSYASI):
            messagebox.showerror("Şablon Bulunamadı",
                                f"Excel şablon dosyası bulunamadı:\n{SABLON_KARSILASTIRMA_DOSYASI}\n"
                                f"Lütfen dosyayı '{SABLONLAR_DIR}' klasörüne yerleştirin.",
                                parent=self.master)
            return

        self.master.update_idletasks()

        try:
            # Gerekli sütunların indekslerini al
            columns_yuzde = self.tree_yuzde['columns']
            try:
                tarih_idx = columns_yuzde.index("Tarih")
                marka_idx = columns_yuzde.index("Cihaz Marka")
                seri_idx = columns_yuzde.index("Cihaz Seri No")
                birim_idx = columns_yuzde.index("Birim/Ünite/Servis Adı")
                hasta_idx = columns_yuzde.index("Hasta Ad Soyad")
                gluko_idx = columns_yuzde.index("Glukometre Sonucu")
                oto_idx = columns_yuzde.index("Oto analizör Sonucu")
                sapma_idx = columns_yuzde.index("% Sapma Oranı")
                degerlendirme_idx = columns_yuzde.index("Değerlendirme Sonucu")
            except ValueError as e:
                messagebox.showerror("Sütun Hatası", f"Yüzde Sapma tablosunda gerekli sütunlardan biri bulunamadı: {e}", parent=self.master)
                return

            all_data_from_tree = []
            for row_id in self.tree_yuzde.get_children():
                values = self.tree_yuzde.item(row_id)['values']
                all_data_from_tree.append(values)

            if not all_data_from_tree:
                messagebox.showinfo("Veri Yok", "Yüzde Sapma tablosunda işlenecek veri bulunamadı.", parent=self.master)
                return

            # Çıktı dosyasını oluşturmak için şablonu kopyala
            timestamp = datetime.now().strftime('%Y.%m.%d_%H-%M')
            output_filename = f"Cihaz_Karsilastirma_Formu_{timestamp}.xlsx"
            output_file_path = os.path.join(EXCEL_OUTPUT_DIR, output_filename)

            shutil.copy(SABLON_KARSILASTIRMA_DOSYASI, output_file_path)
            wb = openpyxl.load_workbook(output_file_path)
            original_sheet_in_output_wb_name = wb.sheetnames[0]
            source_sheet_for_copying = wb[original_sheet_in_output_wb_name]
            chunk_size = 10
            data_chunks = [all_data_from_tree[i:i + chunk_size] for i in range(0, len(all_data_from_tree), chunk_size)]

            active_sheet_modified = False
            for page_index, data_chunk in enumerate(data_chunks):
                current_sheet = None
                if page_index == 0:
                    current_sheet = source_sheet_for_copying
                    current_sheet.title = "Sayfa1"
                    active_sheet_modified = True
                else:
                    current_sheet = wb.copy_worksheet(source_sheet_for_copying)
                    current_sheet.title = f"Sayfa{page_index + 1}"
                for row_num in range(7, 17): # Şablondaki veri başlangıç satırlarına göre ayarlanmış olmalı
                    for col_num in range(1, 9): # Şablondaki sütun sayısına göre
                        current_sheet.cell(row=row_num, column=col_num).value = None

                excel_row = 7 # Excel'e veri yazılacak başlangıç satırı
                for record in data_chunk:
                    tarih = record[tarih_idx]
                    marka = record[marka_idx]
                    seri = record[seri_idx]
                    birim = record[birim_idx]
                    hasta = record[hasta_idx]
                    gluko_sonuc = record[gluko_idx]
                    oto_sonuc = record[oto_idx]
                    sapma_oran_str = record[sapma_idx] # Bu '% Sapma Oranı' string değeridir, örn: "10.50%"
                    # record[degerlendirme_idx] orijinal tek satırlık değerlendirme sonucunu içerir, 
                    # ama biz bunu sapma_oran_str kullanarak yeniden oluşturacağız.

                    # Sütun 1: Tarih
                    current_sheet.cell(row=excel_row, column=1, value=tarih)
                    
                    # Sütun 2: Marka / Seri No
                    marka_seri_cell = current_sheet.cell(row=excel_row, column=2)
                    marka_seri_cell.value = f"{marka}\n{seri}" # Marka ve Seri No aynı Hücre içinde alt alta yaz
                    marka_seri_cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
                    
                    # Sütun 3: Birim/Ünite/Servis Adı
                    current_sheet.cell(row=excel_row, column=3, value=birim)
                    
                    # Sütun 4: Hasta Ad Soyad
                    current_sheet.cell(row=excel_row, column=4, value=hasta)
                    
                    # Sütun 5: Glukometre Sonucu
                    current_sheet.cell(row=excel_row, column=5, value=gluko_sonuc)
                    
                    # Sütun 6: Otoanalizör Sonucu
                    current_sheet.cell(row=excel_row, column=6, value=oto_sonuc)
                    
                    # Sütun 7: % Sapma Oranı
                    current_sheet.cell(row=excel_row, column=7, value=sapma_oran_str)

                    # Sütun 8: Değerlendirme Sonucu (Düzeltilmiş Mantık)
                    excel_formatted_degerlendirme = "" # Başlangıç değeri
                    try:
                        # sapma_oran_str (örn: "12.50%") içinden sayısal değeri al
                        sapma_val = float(sapma_oran_str.replace('%','').strip())
                        
                        if sapma_val > 9.99:
                            excel_formatted_degerlendirme = "□ UYGUN\n✓ UYGUN DEĞİL"
                        else:
                            excel_formatted_degerlendirme = "✓ UYGUN\n□ UYGUN DEĞİL"
                    except ValueError:
                        # Eğer sapma oranı string'i hatalıysa ve float'a çevrilemiyorsa,
                        # ağaçtaki orijinal (tek satırlık) değeri kullan.
                        excel_formatted_degerlendirme = record[degerlendirme_idx] 
                        # print(f"Uyarı: Geçersiz sapma oranı formatı ('{sapma_oran_str}') Excel'e aktarılırken. Satır: {excel_row}. Orijinal değerlendirme kullanılıyor.")
                    
                    cell_col8 = current_sheet.cell(row=excel_row, column=8, value=excel_formatted_degerlendirme)
                    cell_col8.alignment = Alignment(wrap_text=True, vertical='center', horizontal='left')
                    
                    excel_row += 1

            wb.save(output_file_path)

            try:
                if os.name == 'nt': os.startfile(output_file_path)
                elif sys.platform == 'darwin': subprocess.run(['open', output_file_path], check=True)
                else: subprocess.run(['xdg-open', output_file_path], check=True)
            except Exception as e_open:
                print(f"Excel dosyası ({output_file_path}) otomatik olarak açılamadı: {e_open}")

        except ImportError:
            messagebox.showerror("Kütüphane Hatası", "Excel'e aktarma için 'openpyxl' ve 'shutil' kütüphaneleri gereklidir.\nLütfen 'pip install openpyxl' komutu ile yükleyin.", parent=self.master)
        except Exception as e:
            messagebox.showerror("Excel'e Aktarma Hatası", f"Cihaz Karşılaştırma Formu oluşturulurken bir hata oluştu: {e}", parent=self.master)

    def hbtc_formu_olustur(self):
        if not PYTHON_DOCX_AVAILABLE:
            messagebox.showerror("Kütüphane Hatası",
                                 "Word formu oluşturma için 'python-docx' kütüphanesi kurulu değil.\n"
                                 "Lütfen 'pip install python-docx' komutu ile kurun ve programı yeniden başlatın.",
                                 parent=self.master)
            return

        def create_date_selection_dialog():
            dialog = tk.Toplevel(self.master)
            dialog.title("Tarih Aralığı Seçimi")
            dialog.transient(self.master)
            dialog.grab_set()

            # Pencereyi ortala
            dialog.update_idletasks()
            width = 400  
            height = 100
            x = (dialog.winfo_screenwidth() // 2) - (width // 2)
            y = (dialog.winfo_screenheight() // 2) - (height // 2)
            dialog.geometry(f"{width}x{height}+{x}+{y}")
            dialog.resizable(False, False)

            # Tarih girişleri için frame
            date_frame = ttk.Frame(dialog)
            date_frame.pack(pady=10)

            ttk.Label(date_frame, text="Başlangıç Tarihi:").grid(row=0, column=0, padx=5, pady=5)
            start_date_entry = DateEntry(date_frame, width=12, background='darkblue',
                                          foreground='white', borderwidth=2,
                                          date_pattern='dd.mm.yyyy')
            start_date_entry.grid(row=0, column=1, padx=5, pady=5)

            ttk.Label(date_frame, text="Bitiş Tarihi:").grid(row=0, column=2, padx=5, pady=5)
            end_date_entry = DateEntry(date_frame, width=12, background='darkblue',
                                        foreground='white', borderwidth=2,
                                        date_pattern='dd.mm.yyyy')
            end_date_entry.grid(row=0, column=3, padx=5, pady=5)

            def on_ok():
                start_date_str = start_date_entry.get()
                end_date_str = end_date_entry.get()
                try:
                    start_date = datetime.strptime(start_date_str, "%d.%m.%Y").date()
                    end_date = datetime.strptime(end_date_str, "%d.%m.%Y").date()
                    if start_date > end_date:
                        messagebox.showerror("Hata", "Başlangıç tarihi bitiş tarihinden sonra olamaz!", parent=dialog)
                        return
                    dialog.destroy()
                    self.export_to_hbtc(start_date, end_date)
                except ValueError:
                    messagebox.showerror("Hata", "Geçersiz tarih formatı. GG.AA.YYYY kullanın.", parent=dialog)

            ttk.Button(dialog, text="Oluştur", command=on_ok).pack(pady=10)
           
            dialog.focus_set() # Diyalog kutusuna odaklan
            dialog.grab_set()
            dialog.wait_window()

        create_date_selection_dialog()

    def export_to_hbtc(self, start_date, end_date):
        if not self.tree_kalite.get_children():
            messagebox.showinfo("Veri Yok", "Kalite Kontrol tablosu boş. HBTC formu oluşturulacak veri bulunmuyor.", parent=self.master)
            return

        sablon_yolu = os.path.join(SABLONLAR_DIR, HBTC_SABLON_DOSYASI)
        if not os.path.exists(sablon_yolu):
            messagebox.showerror("Şablon Bulunamadı",
                                 f"Şablon dosyası bulunamadı:\n{sablon_yolu}\n"
                                 f"Lütfen '{HBTC_SABLON_DOSYASI}' dosyasını '{SABLONLAR_DIR}' klasörüne yerleştirin.",
                                 parent=self.master)
            return

        self.master.update_idletasks()

        try:
            document = Document(sablon_yolu)
            if not document.tables:
                messagebox.showerror("Şablon Hatası", "Şablon dosyasında doldurulacak tablo bulunamadı.", parent=self.master)
                return

            table = document.tables[0]
            # Tabloyu temizle (varsa mevcut satırları sil)
            for row in table.rows[1:]:
                table._tbl.remove(row._tr)

            col_map = {
                "Tarih": 1, "Cihaz Tipi - Marka": 2, "Cihaz Seri No": 3,
                "L1": 4, "L2": 5, "L3": 6, "Birim/Ünite/Servis Adı": 7
            }
            s_tarih_idx, s_cihaz_adi_idx, s_seri_no_idx = 0, 1, 2
            s_sonuc_idx, s_min_max_idx, s_bolum_idx = 3, 4, 5
            min_max_degerleri = ["36-108 mg/dl", "144-216 mg/dl", "252-396 mg/dl"]

            def set_cell_text(cell, text_lines, is_multiline=False):
                cell.text = ''
                p = cell.paragraphs[0]
                p.text = ''
                if isinstance(text_lines, str): text_lines = [text_lines]
                for i, line_text in enumerate(text_lines):
                    run = p.add_run(line_text)
                    run.font.name = 'Cambria'
                    run.font.size = Pt(12)
                    if is_multiline and i < len(text_lines) - 1:
                        run.add_break(WD_BREAK.LINE)

            for item_id in self.tree_kalite.get_children():
                values = self.tree_kalite.item(item_id)['values']
                tarih_str = values[col_map["Tarih"]]
                try:
                    tarih = datetime.strptime(tarih_str, "%d.%m.%Y").date()
                    if start_date <= tarih <= end_date:
                        row_cells = table.add_row().cells

                        set_cell_text(row_cells[s_tarih_idx], str(values[col_map["Tarih"]]))
                        set_cell_text(row_cells[s_cihaz_adi_idx], str(values[col_map["Cihaz Tipi - Marka"]]))

                        seri_no_ana = str(values[col_map["Cihaz Seri No"]])
                        son_4_hane = seri_no_ana[-4:] if len(seri_no_ana) >= 4 else ""
                        ana_seri = seri_no_ana[:-4] if len(seri_no_ana) > 4 else seri_no_ana
                        set_cell_text(row_cells[s_seri_no_idx], [ana_seri, son_4_hane], is_multiline=True)

                        l1, l2, l3 = str(values[col_map["L1"]]), str(values[col_map["L2"]]), str(values[col_map["L3"]])
                        sonuc_lines = [f"L1 {l1} mg/dl", f"L2 {l2} mg/dl", f"L3 {l3} mg/dl"]
                        set_cell_text(row_cells[s_sonuc_idx], sonuc_lines, is_multiline=True)
                        set_cell_text(row_cells[s_min_max_idx], min_max_degerleri, is_multiline=True)
                        set_cell_text(row_cells[s_bolum_idx], str(values[col_map["Birim/Ünite/Servis Adı"]]))
                except ValueError:
                    print(f"Hatalı tarih formatı: {tarih_str}. Bu kayıt atlandı.")
                    continue

            timestamp = datetime.now().strftime('%Y.%m.%d_%H-%M')
            output_filename = f"HBTC_Kalite_Kontrol_Formu_{timestamp}.docx"
            output_path = os.path.join(HBTC_FORM_OUTPUT_DIR, output_filename)
            document.save(output_path)

            try:
                if os.name == 'nt': os.startfile(output_path)
                elif sys.platform == 'darwin': subprocess.run(['open', output_path], check=True)
                else: subprocess.run(['xdg-open', output_path], check=True)
            except Exception as e_open:
                print(f"Word dosyası ({output_path}) otomatik olarak açılamadı: {e_open}")

        except Exception as e:
            messagebox.showerror("Form Oluşturma Hatası", f"HBTC formu oluşturulurken bir hata oluştu: {e}", parent=self.master)

    def sablondan_excel_e_aktar(self, tree, tablo_adi_kisaltmasi, sablon_dosya_adi):
        if not tree.get_children():
            messagebox.showinfo("Veri Yok", f"{tablo_adi_kisaltmasi} tablosunda aktarılacak veri bulunmuyor.", parent=self.master)
            return

        sablon_yolu = os.path.join(SABLONLAR_DIR, sablon_dosya_adi)
        if not os.path.exists(sablon_yolu):
            messagebox.showerror("Şablon Bulunamadı",
                                 f"Excel şablon dosyası bulunamadı:\n{sablon_yolu}\n"
                                 f"Lütfen '{sablon_dosya_adi}' dosyasını '{SABLONLAR_DIR}' klasörüne yerleştirin.",
                                 parent=self.master)
            return
        
        self.master.update_idletasks()

        try:
            wb = openpyxl.load_workbook(sablon_yolu)
            ws = wb.active
            tree_headers = tree['columns']
            sablon_headers = [str(cell.value).strip() if cell.value is not None else "" for cell in ws[1]]
            
            start_row_excel = 2
            for row_id in tree.get_children():
                values = tree.item(row_id)['values']
                for col_idx, header_title_tree in enumerate(tree_headers):
                    if header_title_tree == "Bir Sonraki Gelinecek Tarih": # Bu sütunu atla
                        continue
                    try:
                        target_col_excel = sablon_headers.index(header_title_tree) + 1
                        cell_value = values[col_idx]
                        current_cell = ws.cell(row=start_row_excel, column=target_col_excel, value=cell_value)
                        current_cell.alignment = Alignment(horizontal='center', vertical='center')

                        if tablo_adi_kisaltmasi == "YuzdeSapma" and header_title_tree == "% Sapma Oranı":
                            try:
                                str_value = str(cell_value).replace('%', '').strip()
                                if str_value:
                                    numeric_value = float(str_value)
                                    if numeric_value > 9.99:
                                        current_cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
                                        current_cell.font = Font(name='Calibri', bold=True, color="FFFFFF")
                                    else:
                                        current_cell.font = Font(name='Calibri')
                                        current_cell.fill = PatternFill(fill_type=None)
                            except ValueError:
                                current_cell.font = Font(name='Calibri')
                                current_cell.fill = PatternFill(fill_type=None)
                        else:
                             current_cell.font = Font(name='Calibri')
                             current_cell.fill = PatternFill(fill_type=None)
                    except ValueError:
                        # print(f"Uyarı: '{header_title_tree}' başlığı şablonda bulunamadı.")
                        continue
                start_row_excel += 1

            if not os.path.exists(EXCEL_OUTPUT_DIR):
                os.makedirs(EXCEL_OUTPUT_DIR)
            timestamp = datetime.now().strftime('%Y.%m.%d_%H.%M')
            output_file_name = f"{tablo_adi_kisaltmasi}_Verileri_{timestamp}.xlsx"
            output_file_path = os.path.join(EXCEL_OUTPUT_DIR, output_file_name)
            wb.save(output_file_path)

            try:
                if os.name == 'nt': os.startfile(output_file_path)
                elif sys.platform == 'darwin': subprocess.run(['open', output_file_path], check=True)
                else: subprocess.run(['xdg-open', output_file_path], check=True)
            except Exception as e_open:
                print(f"Excel dosyası ({output_file_path}) otomatik olarak açılamadı: {e_open}")

        except ImportError:
             messagebox.showerror("Kütüphane Hatası", "Excel'e aktarma için 'openpyxl' kütüphanesi gereklidir.\nLütfen 'pip install openpyxl' komutu ile yükleyin.", parent=self.master)
        except Exception as e:
            messagebox.showerror("Excel'e Aktarma Hatası", f"Veriler Excel'e aktarılırken bir hata oluştu: {e}", parent=self.master)

    def genel_veri_giris_kontrol(self):
        birim = self.cmb_birim.get()
        tip = self.cmb_device_type.get()
        seri = self.cmb_device_serial.get()
        son4 = self.cmb_son4hane.get().strip().upper()

        if not tip:
            messagebox.showwarning("Uyarı", "Cihaz Tipi - Marka seçiniz!")
            self.cmb_device_type.focus_set()
            return False
        if not seri:
            messagebox.showwarning("Uyarı", "Cihaz Seri No seçiniz!")
            self.cmb_device_serial.focus_set()
            return False
        if not (son4 and len(son4) == 4 and son4.isalnum()):
            messagebox.showwarning("Uyarı", "Son 4 Hane 4 karakter ve alfanümerik olmalıdır!")
            self.cmb_son4hane.focus_set()
            return False
        if not birim:
            messagebox.showwarning("Uyarı", "Glukometrenin Geldiği Birim seçiniz!")
            self.cmb_birim.focus_set()
            return False

        baska_birim_atanmis = self.check_device_availability(birim, tip, seri, son4)
        if baska_birim_atanmis:
            messagebox.showerror("Kayıt Çakışması",
                               f"Bu Seri Numaralı Cihaz ({seri}{son4}) \n zaten '{baska_birim_atanmis}' birimine kayıtlıdır.\n"
                               "Girdiğiniz Son 4 Haneyi tekrar kontrol ediniz.")
            return False
        return True

    def ayarla_sonraki_tarih(self, olcum_tarihi_str, gun_ekle):
        try:
            if " " in olcum_tarihi_str:
                olcum_tarihi_dt = datetime.strptime(olcum_tarihi_str.split(" ")[0], "%d.%m.%Y").date()
            else:
                olcum_tarihi_dt = datetime.strptime(olcum_tarihi_str, "%d.%m.%Y").date()

            hedef_tarih = olcum_tarihi_dt + timedelta(days=gun_ekle)
            if hedef_tarih.weekday() == 5:
                hedef_tarih -= timedelta(days=1)
            elif hedef_tarih.weekday() == 6:
                hedef_tarih -= timedelta(days=2)
            return hedef_tarih.strftime("%d.%m.%Y")
        except ValueError:
            print(f"Hatalı tarih formatı: {olcum_tarihi_str}. Bugünkü tarih kullanılıyor.")
            bugun = date.today()
            hedef_tarih = bugun + timedelta(days=gun_ekle)
            if hedef_tarih.weekday() == 5:
                hedef_tarih -= timedelta(days=1)
            elif hedef_tarih.weekday() == 6:
                hedef_tarih -= timedelta(days=2)
            return hedef_tarih.strftime("%d.%m.%Y")

    def tabloya_aktar_kalite(self):
        l_values = self.check_l_entries_valid_for_aktar()
        if l_values is None:
            return

        if not self.genel_veri_giris_kontrol():
            return
        l1, l2, l3 = l_values

        tarih_str = datetime.now().strftime("%d.%m.%Y")
        birim = self.cmb_birim.get()
        cihaz_tipi_marka = self.cmb_device_type.get()
        cihaz_seri_no = self.cmb_device_serial.get()
        son_4_hane = self.cmb_son4hane.get().strip().upper()
        cihaz_seri_no_tam = cihaz_seri_no + son_4_hane  # Seri no ve son 4 haneyi birleştir

        if not self.add_or_update_device_assignment(birim, cihaz_tipi_marka, cihaz_seri_no, son_4_hane):
            return

        sonraki_tarih_str = self.ayarla_sonraki_tarih(tarih_str, 15)

        new_item = self.tree_kalite.insert("", "end", values=(
            self.measurement_no_kalite, tarih_str, cihaz_tipi_marka, cihaz_seri_no_tam,  # Birleştirilmiş seri numarasını kullan
            l1, l2, l3, birim, sonraki_tarih_str
        ))
        self.measurement_no_kalite += 1
        self.txt_l1.delete(0, tk.END); self.txt_l1.config(style="TEntry")
        self.txt_l2.delete(0, tk.END); self.txt_l2.config(style="TEntry")
        self.txt_l3.delete(0, tk.END); self.txt_l3.config(style="TEntry")
        self.tree_kalite.selection_set(new_item)
        self.tree_kalite.see(new_item)
        self.statusbar_guncelle()

    def yuzde_sapma_hesapla_ve_aktar(self):
        if not self.genel_veri_giris_kontrol():
            return
        try:
            glukometre_str = self.txt_glukometre_yuzde.get().strip()
            lab_str = self.txt_lab_yuzde.get().strip()
            if not glukometre_str or not lab_str:
                messagebox.showerror("Hata", "Glukometre ve Laboratuvar ölçüm alanları boş olamaz!")
                return
            glukometre = int(float(glukometre_str))
            lab = int(float(lab_str))
            if glukometre == 0 and lab == 0:
                messagebox.showerror("Hata", "Glukometre ve Laboratuvar ölçümleri aynı anda sıfır olamaz!")
                return
            hasta_ad_soyad = self.txt_hasta_ad_soyad.get().strip()
            if not hasta_ad_soyad:
                messagebox.showerror("Hata", "Hasta Ad Soyad alanı boş olamaz!")
                return
            if glukometre == lab:
                yuzde_sapma = 0.0
            elif min(glukometre, lab) == 0:
                yuzde_sapma = 100.0 if max(glukometre, lab) != 0 else 0.0
            else:
                yuzde_sapma = abs((glukometre - lab) / min(glukometre, lab)) * 100

            tarih_str = datetime.now().strftime("%d.%m.%Y")
            birim = self.cmb_birim.get()
            cihaz_tipi_marka = self.cmb_device_type.get()
            marka = cihaz_tipi_marka.split('-')[-1].strip() if '-' in cihaz_tipi_marka else cihaz_tipi_marka
            cihaz_seri_no = self.cmb_device_serial.get()
            son_4_hane = self.cmb_son4hane.get().strip().upper()
            cihaz_seri_no_tam = cihaz_seri_no + son_4_hane  # Seri no ve son 4 haneyi birleştir
            degerlendirme = "□ UYGUN ✓ UYGUN DEĞİL" if yuzde_sapma > 9.99 else "✓ UYGUN □ UYGUN DEĞİL"

            if not self.add_or_update_device_assignment(birim, cihaz_tipi_marka, cihaz_seri_no, son_4_hane):
                return
            sonraki_tarih_yuzde_str = self.ayarla_sonraki_tarih(tarih_str, 30)
            tags_to_apply = ('high_deviation_tree',) if yuzde_sapma > 9.99 else ()

            new_item = self.tree_yuzde.insert("", "end", values=(self.measurement_no_yuzde, tarih_str, marka, cihaz_seri_no_tam, birim, hasta_ad_soyad, glukometre, lab, f"{yuzde_sapma:.2f}%", degerlendirme, sonraki_tarih_yuzde_str), tags=tags_to_apply)
            self.measurement_no_yuzde += 1
            self.txt_glukometre_yuzde.delete(0, tk.END)
            self.txt_lab_yuzde.delete(0, tk.END)
            self.tree_yuzde.selection_set(new_item)
            self.tree_yuzde.see(new_item)
            self.statusbar_guncelle()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli sayısal değerler girin!")

    def treeview_sort_column(self, tv, col, reverse):
        data_list = []
        for k in tv.get_children(''):
            val = tv.set(k, col)
            data_list.append((val, k))

        try:
            if col == "No":
                data_list.sort(key=lambda t: int(t[0]), reverse=reverse)
            elif col in ["L1", "L2", "L3", "Glukometre Ölçümü", "Lab. Ölçümü"]:
                data_list.sort(key=lambda t: int(float(str(t[0]).replace('%',''))), reverse=reverse)
            elif col == "Yüzde Sapma":
                data_list.sort(key=lambda t: float(str(t[0]).replace('%','')), reverse=reverse)
            elif col in ["Tarih", "Bir Sonraki Gelinecek Tarih"]:
                data_list.sort(key=lambda t: datetime.strptime(str(t[0]), "%d.%m.%Y"), reverse=reverse)
            else:
                data_list.sort(key=lambda t: locale.strxfrm(str(t[0])), reverse=reverse)
        except ValueError as e:
            print(f"Sıralama hatası: {e} - Sütun: {col}, Değer: Unknown")
            data_list.sort(key=lambda t: str(t[0]).upper(), reverse=reverse)
        except NameError:
            data_list.sort(key=lambda t: str(t[0]).upper(), reverse=reverse)

        for index, (_, k) in enumerate(data_list):
            tv.move(k, '', index)
        tv.heading(col, command=lambda: self.treeview_sort_column(tv, col, not reverse))

    def on_double_click(self, event, tree):
        if self.editing_entry:
            self.editing_entry.destroy()
            self.editing_entry = None

        region = tree.identify("region", event.x, event.y)
        if region != "cell":
           
            return
        rowid = tree.identify_row(event.y)
        column_id_str = tree.identify_column(event.x)
        if not rowid or not column_id_str:
            return

        x, y, width, height = tree.bbox(rowid, column_id_str)
        col_index = int(column_id_str.replace("#", "")) - 1

        columns = tree.cget("columns")
        if isinstance(columns, str):
            columns = columns.split()
        if col_index < 0 or col_index >= len(columns):
            return

        column_name = columns[col_index]
        item = tree.item(rowid)
        current_value = item['values'][col_index]

        if column_name == "No":
            messagebox.showwarning("Uyarı", "No sütunu düzenlenemez!")
            return

        self.editing_entry = ttk.Entry(tree)
        self.editing_entry.place(x=x, y=y, width=width, height=height, anchor='nw')
        self.editing_entry.insert(0, current_value)
        self.editing_entry.focus_set()
        self.editing_entry.select_range(0, tk.END)

        def save_edit(event_save=None):
            if not self.editing_entry:
                return
            new_val = self.editing_entry.get().strip()
            current_values_copy = list(item['values'])
            original_tags = tree.item(rowid, 'tags')

            # --- Sütuna özel ilk doğrulamalar ---
            if tree == self.tree_kalite:
                if column_name == "L1":
                    try: val = int(new_val); assert 36 <= val <= 108
                    except (ValueError, AssertionError):
                        messagebox.showerror("Hata", "L1 değeri 36-108 arasında olmalı!", parent=tree)
                        return
                elif column_name == "L2":
                    try: val = int(new_val); assert 144 <= val <= 216
                    except (ValueError, AssertionError):
                        messagebox.showerror("Hata", "L2 değeri 144-216 arasında olmalı!", parent=tree)
                        return
                elif column_name == "L3":
                    try: val = int(new_val); assert 252 <= val <= 396
                    except (ValueError, AssertionError):
                        messagebox.showerror("Hata", "L3 değeri 252-396 arasında olmalı!", parent=tree)
                        return
            
            if column_name in ["Tarih", "Bir Sonraki Gelinecek Tarih"]:
                try:
                    datetime.strptime(new_val, "%d.%m.%Y")
                except ValueError:
                    messagebox.showerror("Hata", f"Geçersiz tarih formatı: {new_val}. GG.AA.YYYY kullanın.", parent=tree)
                    return

            # Düzenlenen hücrenin değerini kopyalanan listede güncelle
            current_values_copy[col_index] = new_val
            new_tags = original_tags

            # --- Yeniden hesaplamalar ve bağımlı güncellemeler ---
            if tree == self.tree_yuzde:
                if column_name in ["Glukometre Sonucu", "Oto analizör Sonucu"]:
                    try:
                        # Düzenlenen değerin bir sayı olduğundan emin ol
                        int(float(new_val)) 
                    except ValueError:
                        messagebox.showerror("Hata", f"{column_name} tamsayı olmalıdır.", parent=tree)
                        return 
                elif column_name == "% Sapma Oranı": # Kullanıcı % Sapma'yı doğrudan düzenlerse doğrulama
                    try:
                        if not new_val.endswith('%'): raise ValueError("Yüzde işareti eksik.")
                        float(new_val.replace('%','')) # Sayısal kısmın geçerli olup olmadığını kontrol et
                    except ValueError as e:
                        messagebox.showerror("Hata", f"Yüzde Sapma formatı hatalı (örn: 10.50%): {e}", parent=tree)
                        return

                # Glukometre veya Lab sonuçları değiştiyse yeniden hesaplama
                if column_name in ["Glukometre Sonucu", "Oto analizör Sonucu"]:
                    try:
                        gluk_idx = columns.index("Glukometre Sonucu")
                        lab_idx = columns.index("Oto analizör Sonucu")
                        sapma_idx = columns.index("% Sapma Oranı")
                        degerlendirme_idx = columns.index("Değerlendirme Sonucu") # EKLENDİ

                        # current_values_copy listesi zaten yeni değeri içeriyor
                        glukometre_str_val = current_values_copy[gluk_idx]
                        lab_str_val = current_values_copy[lab_idx]
                        
                        # Değerlerin ayrıştırılabilir tamsayılar olduğundan emin ol
                        glukometre = int(float(glukometre_str_val))
                        lab = int(float(lab_str_val))

                        if glukometre == lab:
                            yuzde_sapma = 0.0
                        elif min(glukometre, lab) == 0: # Sıfıra bölme hatasını önle
                            yuzde_sapma = 100.0 if max(glukometre, lab) != 0 else 0.0
                        else:
                            yuzde_sapma = abs((glukometre - lab) / min(glukometre, lab)) * 100
                        
                        current_values_copy[sapma_idx] = f"{yuzde_sapma:.2f}%"
                        degerlendirme = "□ UYGUN ✓ UYGUN DEĞİL" if yuzde_sapma > 9.99 else "✓ UYGUN □ UYGUN DEĞİL"
                        current_values_copy[degerlendirme_idx] = degerlendirme
                        
                        new_tags = ('high_deviation_tree',) if yuzde_sapma > 9.99 else ()
                    
                    except ValueError:
                        # Bu hata, gluk_idx/lab_idx değerlerinin dönüştürülemediği anlamına gelir.
                        # Bu durum, yukarıdaki sütunlara özel doğrulama tarafından yakalanmalıydı.
                        messagebox.showerror("Hata", "Yüzde sapma hesaplanamadı! Glukometre/Lab değerlerini kontrol edin.", parent=tree)
                        return # Önemli: Dönüşüm başarısız olursa işlemeyi durdur
            
            # "Tarih" değiştirildiyse "Bir Sonraki Gelinecek Tarih"i güncelle
            if column_name == "Tarih":
                try:
                    sonraki_tarih_idx = columns.index("Bir Sonraki Gelinecek Tarih")
                    gun_ekle = 15 if tree == self.tree_kalite else 30 # Tabloya göre gün ekle
                    current_values_copy[sonraki_tarih_idx] = self.ayarla_sonraki_tarih(new_val, gun_ekle)
                except ValueError: # Sütun bulunamazsa (örn. farklı bir tabloda)
                    pass

            # Satırı tüm değişikliklerle son olarak güncelle
            tree.item(rowid, values=tuple(current_values_copy), tags=new_tags) 
            
            if self.editing_entry:
                self.editing_entry.destroy()
            self.editing_entry = None

        self.editing_entry.bind("<Return>", save_edit)
        self.editing_entry.bind("<FocusOut>", save_edit)
        self.editing_entry.bind("<Escape>", lambda e: (self.editing_entry.destroy(), setattr(self, 'editing_entry', None)))

    def show_context_menu(self, event, tree):
        rowid = tree.identify_row(event.y)
        if rowid:
            if not tree.selection():
                tree.selection_set(rowid)
            elif rowid not in tree.selection():
                tree.selection_set(rowid)

            if tree == self.tree_kalite:
                self.context_menu_kalite.post(event.x_root, event.y_root)
            elif tree == self.tree_yuzde:
                self.context_menu_yuzde.post(event.x_root, event.y_root)

    def satir_sil(self, tree, is_kalite_table):
        selected_items = tree.selection()
        if not selected_items:
            messagebox.showwarning("Uyarı", "Silmek için bir satır seçiniz!")
            return
        msg = "Seçilen satırı silmek istediğinize emin misiniz?"
        if len(selected_items) > 1:
            msg = f"{len(selected_items)} adet seçili satırı silmek istediğinize emin misiniz?"
        if messagebox.askokcancel("Onay", msg):
            for sel_item in selected_items:
                tree.delete(sel_item)
            self.guncelle_no_sutunu(tree, is_kalite_table)
            self.statusbar_guncelle()

    def guncelle_no_sutunu(self, tree, is_kalite_table):
        all_items = tree.get_children('')
        for idx, item_id in enumerate(all_items, start=1):
            current_values = list(tree.item(item_id, "values"))
            current_values[0] = idx
            tree.item(item_id, values=tuple(current_values))

        if is_kalite_table:
            self.measurement_no_kalite = len(all_items) + 1
        else:
            self.measurement_no_yuzde = len(all_items) + 1

    def tablolari_temizle(self):
        if messagebox.askokcancel("Onay", "Tüm tablolardaki verileri silmek istediğinize emin misiniz? " \
        "Bu işlem yedeklenmiş verileri etkilemez, sadece mevcut görünümü temizler."):
            for item in self.tree_kalite.get_children():
                self.tree_kalite.delete(item)
            for item in self.tree_yuzde.get_children():
                self.tree_yuzde.delete(item)
            self.measurement_no_kalite = 1
            self.measurement_no_yuzde = 1
            self.tables_cleared_this_session = True
            self.statusbar_guncelle()

    def clear_backup_folder(self):
        msg = "BU İŞLEM TAMAMEN GÜVENLİDİR:\n\n" \
              "- Son 2 yedekleme korunacak, Sadece daha eski yedeklemeler silinecektir \n" \
              "\n" \
              "Devam etmek istiyor musunuz?"

        if not messagebox.askokcancel("Yedeklenmiş Verileri Temizleme", msg, parent=self.master):
            return

        try:
            backup_path = BACKUP_DIR
            if not os.path.exists(backup_path):
                messagebox.showinfo("Bilgi", "Yedeklenmiş Veriler klasörü bulunamadı.", parent=self.master)
                return

            all_files = glob.glob(os.path.join(backup_path, "*.csv"))
            all_files = [f for f in all_files if not os.path.basename(f).startswith('~$')]

            kalite_yedekleri = sorted([f for f in all_files if "Kalite_Kontrol_Olcumleri_Yedek_" in os.path.basename(f)], key=os.path.getctime)          
            yuzde_yedekleri = sorted([f for f in all_files if "Yuzde_Sapma_Olcumleri_Yedek_" in os.path.basename(f)], key=os.path.getctime)

            files_to_delete = []
            if len(kalite_yedekleri) > 2:
                files_to_delete.extend(kalite_yedekleri[:-2])
            if len(yuzde_yedekleri) > 2:
                files_to_delete.extend(yuzde_yedekleri[:-2])

            if not files_to_delete:
                messagebox.showinfo("Bilgi", "Silinecek daha eski yedek bulunamadı (Her tablo için son 2 yedek dosyası korunuyor).", parent=self.master)
                return

            deleted_count = 0
            for file_path_del in files_to_delete:
                try:
                    os.remove(file_path_del)
                    print(f"Silindi: {os.path.basename(file_path_del)}")
                    deleted_count +=1
                except Exception as e_del:
                    print(f"Silme hatası: {os.path.basename(file_path_del)} - {e_del}")
            if deleted_count > 0:
                messagebox.showinfo("Bilgi", f"Her tablo için son 2 yedekleme saklandı. {deleted_count} adet eski yedek dosyası silindi. ", parent=self.master)
            else:
                messagebox.showinfo("Bilgi", "Silinecek ek yedek bulunamadı.", parent=self.master)
        except Exception as e:
            messagebox.showerror("Hata", f"Yedeklenmis Veriler klasörü temizlenirken bir hata oluştu: {e}", parent=self.master)

    def _tarih_farki_hesapla(self, hedef_tarih_str):
        try:
            hedef_tarih = datetime.strptime(hedef_tarih_str, "%d.%m.%Y").date()
            bugun = date.today()
            fark = (hedef_tarih - bugun).days
            return fark
        except ValueError:
            return None

    def _topla_olcum_verileri(self, durum_tipi):
            bulunan_olcumler = []
            bugun = date.today()
            
            # Adım 1: Her bir benzersiz cihaz için en son ölçüm kaydının detaylarını bul.
            # Anahtar: (birim_adi, cihaz_seri_no)
            # Değer: (en_son_olcum_tarihi_dt, gelmesi_gereken_tarih_str, olcum_tipi_str, birim_adi, cihaz_seri_no)
            latest_measurement_details = {}

            def find_latest_measurements(tree, olcum_tipi_str, birim_idx, seri_idx, olcum_tarih_idx_tree, gelmesi_gereken_tarih_idx_tree):
                for item_id in tree.get_children():
                    values = tree.item(item_id, 'values')
                    try:
                        if len(values) <= max(birim_idx, seri_idx, olcum_tarih_idx_tree, gelmesi_gereken_tarih_idx_tree):
                            # print(f"Uyarı: '{olcum_tipi_str}' tablosunda eksik sütunlu veri: {values}")
                            continue

                        birim_adi = values[birim_idx]
                        cihaz_seri_no = values[seri_idx] # Bu zaten tam seri no (ana + son4hane) olmalı
                        olcum_tarihi_str = values[olcum_tarih_idx_tree]
                        gelmesi_gereken_tarih_str_val = values[gelmesi_gereken_tarih_idx_tree]

                        olcum_tarihi_dt = datetime.strptime(olcum_tarihi_str, "%d.%m.%Y").date()
                        cihaz_anahtar = (birim_adi, cihaz_seri_no)

                        if cihaz_anahtar not in latest_measurement_details or \
                        olcum_tarihi_dt > latest_measurement_details[cihaz_anahtar][0]:
                            latest_measurement_details[cihaz_anahtar] = (
                                olcum_tarihi_dt,
                                gelmesi_gereken_tarih_str_val,
                                olcum_tipi_str,
                                birim_adi, 
                                cihaz_seri_no
                            )
                    except (ValueError, IndexError) as e:
                        print(f"Hata (_topla_olcum_verileri - Adım 1 - {olcum_tipi_str}): {e} - Veri: {values}")
                        continue
            
            # Kalite Kontrol Ağacı Sütun İndeksleri:
            # "No"(0), "Tarih"(1), ..., "Cihaz Seri No"(3), ..., "Birim/Ünite/Servis Adı"(7), "Bir Sonraki Gelinecek Tarih"(8)
            find_latest_measurements(self.tree_kalite, "Kalite Kontrol", 
                                    birim_idx=7, seri_idx=3, olcum_tarih_idx_tree=1, gelmesi_gereken_tarih_idx_tree=8)

            # Yüzde Sapma Ağacı Sütun İndeksleri:
            # "No"(0), "Tarih"(1), ..., "Cihaz Seri No"(3), "Birim/Ünite/Servis Adı"(4), ..., "Bir Sonraki Gelinecek Tarih"(10)
            find_latest_measurements(self.tree_yuzde, "Yüzde Sapma",
                                    birim_idx=4, seri_idx=3, olcum_tarih_idx_tree=1, gelmesi_gereken_tarih_idx_tree=10)

            # Adım 2: Bulunan en son ölçümlere göre durumu değerlendir.
            sira_no_counter = 1
            temp_olcumler = []

            for cihaz_anahtar, data in latest_measurement_details.items():
                # data = (en_son_olcum_tarihi_dt, gelmesi_gereken_tarih_str, olcum_tipi_str, birim_adi, cihaz_seri_no)
                gelmesi_gereken_tarih_str_cihazin = data[1]
                olcum_tipi_cihazin = data[2]
                birim_adi_cihazin = data[3]
                cihaz_seri_no_cihazin = data[4]

                try:
                    gelmesi_gereken_tarih_dt = datetime.strptime(gelmesi_gereken_tarih_str_cihazin, "%d.%m.%Y").date()
                    fark_gun = (gelmesi_gereken_tarih_dt - bugun).days

                    if durum_tipi == "gecen" and fark_gun < 0:
                        gecen_gun_str = f"{-fark_gun} gün geçti"
                        kayit = (0, olcum_tipi_cihazin, birim_adi_cihazin, cihaz_seri_no_cihazin,
                                gelmesi_gereken_tarih_str_cihazin, gecen_gun_str, gelmesi_gereken_tarih_dt) # Sıralama için dt ekle
                        temp_olcumler.append(kayit)
                    
                    elif durum_tipi == "yaklasan" and 0 <= fark_gun <= 2: # Yaklaşanlar için 0-2 gün aralığı
                        kalan_gun_str = f"{fark_gun} gün kaldı"
                        if fark_gun == 0: 
                            kalan_gun_str = "Bugün"
                        kayit = (0, olcum_tipi_cihazin, birim_adi_cihazin, cihaz_seri_no_cihazin,
                                gelmesi_gereken_tarih_str_cihazin, kalan_gun_str, gelmesi_gereken_tarih_dt) # Sıralama için dt ekle
                        temp_olcumler.append(kayit)
                
                except ValueError as e:
                    print(f"Hata (_topla_olcum_verileri - Adım 2 - {durum_tipi}): {e} - Tarih Str: {gelmesi_gereken_tarih_str_cihazin}")
                    continue
            
            # Geçerlilik tarihine göre sırala (en eski geçerlilik tarihi en üstte)
            temp_olcumler.sort(key=lambda x: x[6]) 

            # Sıra numaralarını ata ve son listeyi oluştur
            for kayit_data in temp_olcumler:
                bulunan_olcumler.append((sira_no_counter,) + kayit_data[1:6]) # dt object'i (kayit_data[6]) son listeye ekleme
                sira_no_counter += 1
                
            return bulunan_olcumler

    def _goster_durum_penceresi(self, baslik, olcum_listesi):
        if not olcum_listesi:
            messagebox.showinfo(baslik, "Bu kriteri karşılayan ölçüm bulunmuyor.", parent=self.master)
            return

        top = tk.Toplevel(self.master)
        top.title(baslik)
        top.geometry("800x400")

        columns = ("No", "Ölçüm Tipi", "Cihazın Geldiği Birim", "Cihaz Seri No", "Gelmesi Gereken Tarih", "Durum")
        tree_durum = ttk.Treeview(top, columns=columns, show="headings")
        tree_durum.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        vsb = ttk.Scrollbar(top, orient="vertical", command=tree_durum.yview)
        vsb.pack(side="right", fill="y", pady=(10,10))
        tree_durum.configure(yscrollcommand=vsb.set)

        widths = [40, 120, 200, 150, 130, 120]
        for i, col_name in enumerate(columns):
            tree_durum.heading(col_name, text=col_name, 
                         command=lambda c=col_name: self.treeview_sort_column(tree_durum, c, False))
            tree_durum.column(col_name, width=widths[i], anchor=tk.W if i in [2,3] else tk.CENTER, minwidth=widths[i])

        for olcum_verisi in olcum_listesi:
            tree_durum.insert("", "end", values=olcum_verisi)

        btn_kapat = ttk.Button(top, text="Kapat", command=top.destroy)
        btn_kapat.pack(pady=10)

        top.update_idletasks()
        width = top.winfo_width()
        height = top.winfo_height()
        x = (top.winfo_screenwidth() // 2) - (width // 2)
        y = (top.winfo_screenheight() // 2) - (height // 2)
        top.geometry(f'{width}x{height}+{x}+{y}')
        top.focus_set()
        top.wait_window()

    def goster_gunu_gecen_olcumler(self):
        gecen_olcumler = self._topla_olcum_verileri("gecen")
        self._goster_durum_penceresi("Günü Geçen Ölçümler", gecen_olcumler)

    def goster_gunu_yaklasan_olcumler(self):
        yaklasan_olcumler = self._topla_olcum_verileri("yaklasan")
        self._goster_durum_penceresi("Günü Yaklaşan Ölçümler (Son 2 Gün)", yaklasan_olcumler)

    def save_data_to_timestamped_csv(self):
        self.master.update_idletasks()
        try:
            timestamp = datetime.now().strftime("%Y.%m.%d_%H.%M")
            kalite_has_data = bool(self.tree_kalite.get_children())
            yuzde_has_data = bool(self.tree_yuzde.get_children())

            if kalite_has_data:
                kalite_filename_ts = os.path.join(BACKUP_DIR, f"Kalite_Kontrol_Olcumleri_Yedek_{timestamp}.csv")
                with open(kalite_filename_ts, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(self.tree_kalite['columns'])
                    for row_id in self.tree_kalite.get_children():
                        writer.writerow(self.tree_kalite.item(row_id)['values'])
                print(f"Kalite Kontrol verileri {os.path.basename(kalite_filename_ts)} dosyasına yedeklendi.")

            if yuzde_has_data:
                yuzde_filename_ts = os.path.join(BACKUP_DIR, f"Yuzde_Sapma_Olcumleri_Yedek_{timestamp}.csv")
                with open(yuzde_filename_ts, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(self.tree_yuzde['columns'])
                    for row_id in self.tree_yuzde.get_children():
                        writer.writerow(self.tree_yuzde.item(row_id)['values'])
                print(f"Yüzde Sapma verileri {os.path.basename(yuzde_filename_ts)} dosyasına yedeklendi.")

            if not kalite_has_data and not yuzde_has_data:
                 print("Kaydedilecek veri bulunmadığı için yedekleme yapılmadı.")
        except Exception as e:
            messagebox.showerror(".CSV Kaydetme Hatası", f"Veriler .CSV dosyasına kaydedilirken bir hata oluştu:\n{e}", parent=self.master)


    def is_iptv_db_empty(self):
        """iptv_kanallar tablosunun boş olup olmadığını kontrol eder."""
        try:
            conn = sqlite3.connect(VERITABANI_DOSYASI)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM iptv_kanallar")
            count = cursor.fetchone()[0]
            conn.close()
            return count == 0
        except sqlite3.Error as e:
            print(f"IPTV veritabanı kontrol hatası: {e}")
            return True  # Hata durumunda varsayılan olarak boş kabul et

    def iptv_playlist_yukle_diyalog(self):
        # Eğer IPTV sekmesi zaten varsa o sekmeye geç
        for i in range(self.notebook.index("end")):
            if self.notebook.tab(i, "text") == "IPTV":
                self.notebook.select(i)
                return

        # Yoksa sekmeyi oluştur
        from Modüller import iptv_modul
        iptv_modul.create_iptv_tab(self, self.notebook, self.iptv_channels)
        # Sekmeyi seç
        for i in range(self.notebook.index("end")):
            if self.notebook.tab(i, "text") == "IPTV":
                self.notebook.select(i)
                return

    def load_iptv_playlist(self, url):
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            lines = response.text.splitlines()
            channels = []
            name, stream_url = None, None

            for line in lines:
                if line.startswith("#EXTINF"):
                    name = line.split(",")[-1].strip()
                elif line and not line.startswith("#"):
                    stream_url = line.strip()
                    if name and stream_url:
                        channels.append((name, stream_url))
                        name, stream_url = None, None

        except Exception as e:
            messagebox.showerror("Hata", f"Playlist yüklenemedi:\n{e}", parent=self.master)

############## ANA MENÜ #######################
    def create_menu(self):
        menubar = tk.Menu(self.master)
        menu_dosya = tk.Menu(menubar, tearoff=0)
        menu_dosya.add_command(label="Verileri Yedekten Geri Yükle...", command=self.manuel_yedek_yukle)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Tüm Tabloları Temizle", command=self.tablolari_temizle)
        menu_dosya.add_command(label="Yedeklenmiş Veriler Klasörünü Temizle", command=self.clear_backup_folder)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Hakkında", command=self.launch_hakkinda)
        menu_dosya.add_separator()
        menu_dosya.add_command(label="Çıkış", command=self.kaydet_ve_cikis_yap)
        menubar.add_cascade(label="Dosya", menu=menu_dosya)

        menu_cihaz_takip = tk.Menu(menubar, tearoff=0)
        menu_cihaz_takip.add_command(label="Cihaz Ara", command=self.open_cihaz_arama_dialog, accelerator="Ctrl+F")
        menu_cihaz_takip.add_command(label="Cihaz Ekle/Sil", command=self.open_cihaz_ekle_sil_dialog)
        menu_cihaz_takip.add_command(label="Cihaz Değişim Formu Oluştur")
        menubar.add_cascade(label="Cihaz Takibi", menu=menu_cihaz_takip)

        menu_kalite_kontrol = tk.Menu(menubar, tearoff=0)
        menu_kalite_kontrol.add_command(label="HBTC Formu Oluştur", command=self.hbtc_formu_olustur)
        menu_kalite_kontrol.add_command(label="Verileri Excel'e Aktar", command=lambda: self.sablondan_excel_e_aktar(self.tree_kalite, "KaliteKontrol", KALITE_KONTROL_EXCEL_SABLON_DOSYASI))
        menubar.add_cascade(label="Kalite Kontrol Menüsü", menu=menu_kalite_kontrol)

        menu_yuzde_sapma = tk.Menu(menubar, tearoff=0)
        menu_yuzde_sapma.add_command(label="Cihaz Karşılaştırma Sonuç Formu Oluştur", command=self.cihaz_karsilastirma_formu_olustur)
        menu_yuzde_sapma.add_command(label="Verileri Excel'e Aktar", command=lambda: self.sablondan_excel_e_aktar(self.tree_yuzde, "YuzdeSapma", YUZDE_SAPMA_EXCEL_SABLON_DOSYASI))
        menubar.add_cascade(label="Yüzde Sapma Menüsü", menu=menu_yuzde_sapma)

        menu_durum_tespiti = tk.Menu(menubar, tearoff=0)
        menu_durum_tespiti.add_command(label="Günü Geçen Ölçümler", command=self.goster_gunu_gecen_olcumler, accelerator="Ctrl+G")
        menu_durum_tespiti.add_command(label="Günü Yaklaşan Ölçümler", command=self.goster_gunu_yaklasan_olcumler, accelerator="Ctrl+Y")
        menubar.add_cascade(label="Durum Tespiti", menu=menu_durum_tespiti)

        menu_ekstralar = tk.Menu(menubar, tearoff=0)
        menu_ekstralar.add_command(label="Ekran Görüntüsü Al (1 sn sonra)", command=self.take_screenshot, accelerator="Ctrl+F8")
        menu_ekstralar.add_command(label="IPTV Aktif Et", command=self.iptv_playlist_yukle_diyalog)
        menu_ekstralar.add_command(label="Hesap Makinesi", command=self.open_calculator, accelerator="Ctrl+H")
        menu_ekstralar.add_command(label="Notepad", command=self.launch_notepad, accelerator="Ctrl+N")        
        menu_ekstralar.add_command(label="Takvim - Ajanda", command=self.launch_calendar, accelerator="Ctrl+T")
        menu_ekstralar.add_command(label="Harita", command=self.open_map, accelerator="Ctrl+M")
        menu_ekstralar.add_command(label="Sağlık Hesaplamaları", command=self.launch_bmi_calculator)
        menu_ekstralar.add_command(label="Excel / CSV Görüntüleyici", command=self.launch_excel_csv_viewer)
        menu_ekstralar.add_command(label="Flappy Bird Oyunu", command=self.launch_flappy_bird)
        menu_ekstralar.add_command(label="Yılan Oyunu", command=self.snake_oyunu)
        menu_ekstralar.add_command(label="Sayı Tahmin Oyunu", command=self.start_cows_bulls_game)
        menu_ekstralar.add_command(label="2048 Oyunu", command=self.launch_game_2048)
        menu_ekstralar.add_command(label="Memory Puzzle Oyunu", command=self.launch_memory_puzzle)
        menubar.add_cascade(label="Ekstra Özellikler", menu=menu_ekstralar)
        self.master.config(menu=menubar)
        self.master.bind('<Control-h>', lambda event: self.open_calculator())
        self.master.bind('<Control-t>', lambda event: self.launch_calendar())
        self.master.bind('<Control-f>', lambda event: self.open_cihaz_arama_dialog())
        self.master.bind('<Control-g>', lambda event: self.goster_gunu_gecen_olcumler())
        self.master.bind('<Control-m>', lambda event: self.open_map())
        self.master.bind('<Control-n>', lambda event: self.launch_notepad())
        self.master.bind('<Control-F8>', lambda event: self.take_screenshot())
        self.master.bind('<Control-y>', lambda event: self.goster_gunu_yaklasan_olcumler())
###################################
####################################KODUN AYRILABİLİR KISMI####################################
#    def launch_paint(self):
#        try:
#            script_dir = os.path.dirname(os.path.abspath(__file__))
#            paint_path = os.path.join(script_dir, "Modüller\\Paint\\paint.py")
#            if not os.path.exists(paint_path):
#                messagebox.showerror("Hata", f"'paint.py' dosyası bulunamadı:\n{paint_path}")
#                return
#            subprocess.Popen([sys.executable, paint_path])
#        except Exception as e:
#            messagebox.showerror("Hata", f"Paint uygulaması başlatılamadı:\n{e}")

    def launch_game_2048(self):
        subprocess.Popen([sys.executable, "Modüller\\game_2048.py"])

    def start_cows_bulls_game(self):
        CowsAndBullsGame(self.master)

    def launch_flappy_bird(self):
        subprocess.Popen([sys.executable, "Modüller\\Flappy Bird\\main.py"])

    def snake_oyunu(self):
        snake_game.run_snake_game()

    def launch_memory_puzzle(self):
        subprocess.Popen([sys.executable, "Modüller\\Memory Puzzle\\game.py"])

    def launch_hakkinda(self):
        show_about(self.master)

    def launch_notepad(self):
        try:
            if os.name == 'nt': subprocess.run("notepad.exe", shell=True, check=True)
        except (FileNotFoundError, subprocess.CalledProcessError) as e:
            messagebox.showerror("Hata", f"Notepad açılamadı: {e}", parent=self.master)

    def launch_excel_csv_viewer(self):
        subprocess.Popen([sys.executable, "Modüller\\excel_csv_viewer.py"])

    def launch_bmi_calculator(self):
        subprocess.Popen([sys.executable, "Modüller\\bmi_calculator.py"])

    def launch_calendar(self, event=None):
        try:
            ajanda.show_agenda_ui(self.master, VERITABANI_DOSYASI)
        except Exception as e:
            print(f"Ajanda açılamadı: {e}")
            messagebox.showerror("Ajanda Hatası", f"Ajanda arayüzü açılamadı:\n{e}", parent=self.master)

    def initialize_agenda_module(self):
        try:
            ajanda.init_agenda_db(VERITABANI_DOSYASI)
            ajanda.show_startup_alerts(self.master, VERITABANI_DOSYASI) # Başlangıçta Hatırlatma penceresi göster (Eğer kayıt mevcutsa)
        except Exception as e:
            print(f"Ajanda modülü başlatılırken hata: {e}")
            messagebox.showerror("Ajanda Hatası", f"Ajanda modülü başlatılamadı:\n{e}", parent=self.master)

    def take_screenshot(self):
        try:
            import time
            from PIL import ImageGrab
            from datetime import datetime
            time.sleep(1)  # 1 saniye bekledikten sonra
            now = datetime.now().strftime("%Y.%m.%d_%H-%M-%S")
            filename = f"screenshot_{now}.png"
            ImageGrab.grab().save(filename)
            messagebox.showinfo("Ekran Görüntüsü", f"Ekran görüntüsü kaydedildi:\n{filename}")
        except Exception as e:
            messagebox.showerror("Hata", f"Ekran görüntüsü alınamadı:\n{e}")
            self.master.deiconify()
   
    def open_calculator(self):
        try:
            if os.name == 'nt': subprocess.run("calc.exe", shell=True, check=True)
            elif sys.platform == 'darwin': subprocess.run(["open", "-a", "Calculator"], check=True)
            else:
                found = False
                for calc_cmd in ["gnome-calculator", "kcalc", "xcalc"]:
                    try: subprocess.run([calc_cmd], check=True); found = True; break
                    except FileNotFoundError: continue
                if not found: messagebox.showwarning("Hesap Makinesi", "Standart hesap makinesi komutları bulunamadı.", parent=self.master)
        except (FileNotFoundError, subprocess.CalledProcessError) as e:
            messagebox.showerror("Hata", f"Hesap makinesi açılamadı: {e}", parent=self.master)

    def open_map(self):
        self.map_viewer.open_map()
############################################################################################
    def kaydet_ve_cikis_yap(self):
        if hasattr(self, '_is_closing') and self._is_closing:
            return
        self._is_closing = True

        try:
            # Pencere durumunu ayarlara kaydet
            if self.master.winfo_exists():
                is_maximized = "1" if self.master.state() == 'zoomed' else "0"
                self.program_ayari_kaydet("window_maximized", is_maximized)
                if is_maximized == "0":
                    window_geometry = self.master.geometry()
                    self.program_ayari_kaydet("window_geometry", window_geometry)
                else:
                    self.program_ayari_kaydet("window_geometry", "")

            # Radyo istasyonu ve ses seviyesini ayarlara kaydet
            last_station_name = self.cmb_radyo.get()
            if self.radio_process and self.radio_process.poll() is None:
                self.stop_radio()
            self.radio_process = None
            if last_station_name:
                self.program_ayari_kaydet("last_radio_station", last_station_name)
            self.program_ayari_kaydet("last_radio_volume", str(self.radio_volume.get()))

            # Mevcut verileri kontrol et
            kalite_has_data = bool(self.tree_kalite.get_children())
            yuzde_has_data = bool(self.tree_yuzde.get_children())
            tables_have_data = kalite_has_data or yuzde_has_data
            if not tables_have_data and self.tables_cleared_this_session:
                if self.master.winfo_exists() and messagebox.askokcancel("Çıkış Onayı", 
                                                                         "Tablolar temizlendi ve yeni veri girilmedi.\nÇıkmak istediğinize emin misiniz?",
                                                                           parent=self.master):
                    self.master.destroy()
                return
            elif not tables_have_data and not self.tables_cleared_this_session:
                self.master.destroy()
                return
            else:
                if self.master.winfo_exists() and messagebox.askokcancel("Yedekleme Yapalım mı?", 
                                                                         "TAMAM: Mevcut verileri yeni zaman damgalı dosyalara yedekleyelim ve çıkalım." \
                                                                         "\n\nİPTAL: Yedekleme yapmadan çıkalım",
                                                                           parent=self.master):
                    self.save_data_to_timestamped_csv()
                    self.analog_saat_instance.stop()
                    self.master.destroy()
        finally:
            self.master.after(200, self.master.destroy)  # Güvenli kapanış

if __name__ == '__main__':
    root = tk.Tk()
    app = MainWindow(root)
    root.mainloop()
